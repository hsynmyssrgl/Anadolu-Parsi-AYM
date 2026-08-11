from pathlib import Path
import json, collections, hashlib, textwrap
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/mnt/data/Anadolu_Parsi_Aile_Yasam_Merkezi_Bronze_04.08.2026.28_Kaynak')
OUT=ROOT/'docs/current/MASTER_PROJE_DOKUMANTASYONU_BRONZE_04.08.2026.28.docx'

rules=json.loads((ROOT/'config/canonical-rule-registry.json').read_text(encoding='utf-8'))
constitution=json.loads((ROOT/'config/project-constitution.json').read_text(encoding='utf-8'))
scope=json.loads((ROOT/'config/accepted-scope-registry.json').read_text(encoding='utf-8'))
decisions=json.loads((ROOT/'config/user-decision-ledger.json').read_text(encoding='utf-8'))
release=json.loads((ROOT/'config/release-ledger.json').read_text(encoding='utf-8'))
governance=json.loads((ROOT/'config/active-governance-ledger.json').read_text(encoding='utf-8'))
progress=json.loads((ROOT/'artifacts/reports/PROJECT_PROGRESS_04.08.2026.28.json').read_text(encoding='utf-8'))
delivery=json.loads((ROOT/'artifacts/reports/DELIVERY_STATUS_04.08.2026.28.json').read_text(encoding='utf-8'))
active_docs=json.loads((ROOT/'config/active-document-set.json').read_text(encoding='utf-8'))

reqs=scope['requirements']
status_counts=collections.Counter(r['status'] for r in reqs)
priority_counts=collections.Counter(r['priority'] for r in reqs)
area_counts=collections.Counter(r['area'] for r in reqs)

TEAL='0E6666'; GOLD='A87321'; DARK='163535'; LIGHT='EAF4F2'; PALE='F7F4ED'; RED='A13632'; GRAY='5C666A'; WHITE='FFFFFF'; BLACK='1D2728'

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr()
    shd=tcPr.find(qn('w:shd'))
    if shd is None:
        shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell(cell, text, bold=False, color=BLACK, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text=''
    p=cell.paragraphs[0]; p.alignment=align; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(str(text)); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def table(rows, widths=None, header=True):
    t=doc.add_table(rows=1 if header else 0, cols=len(rows[0]))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    if header:
        set_repeat_table_header(t.rows[0])
        for c,txt in zip(t.rows[0].cells,rows[0]):
            set_cell(c,txt,True,WHITE,8.2,WD_ALIGN_PARAGRAPH.CENTER); shade(c,TEAL); set_cell_margins(c)
        body=rows[1:]
    else: body=rows
    for i,row in enumerate(body):
        cells=t.add_row().cells
        for j,(c,txt) in enumerate(zip(cells,row)):
            set_cell(c,txt,False,BLACK,8.1); shade(c, WHITE if i%2==0 else 'F1F6F5'); set_cell_margins(c)
    if widths:
        for row in t.rows:
            for c,w in zip(row.cells,widths): c.width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def bullet(text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2')
    p.paragraph_format.space_after=Pt(2); p.add_run(text)
    return p

def heading(text, level=1):
    p=doc.add_heading(text, level=level); p.paragraph_format.keep_with_next=True; return p

def callout(title, text, fill=PALE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(t.rows[0])
    c=t.cell(0,0); shade(c,fill); set_cell_margins(c,160,180,160,180)
    c.text=''; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3)
    r=p.add_run(title); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(TEAL)
    p=c.add_paragraph(text); p.paragraph_format.space_after=Pt(0)
    for r in p.runs: r.font.name='Arial'; r.font.size=Pt(9)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

doc=Document()
sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.72); sec.right_margin=Inches(.72)
sec.header_distance=Inches(.25); sec.footer_distance=Inches(.25)
styles=doc.styles
styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(9.5); styles['Normal'].font.color.rgb=RGBColor.from_string(BLACK)
styles['Normal'].paragraph_format.space_after=Pt(5); styles['Normal'].paragraph_format.line_spacing=1.08
for style_name,size,color in [('Title',26,TEAL),('Heading 1',16,TEAL),('Heading 2',12,GOLD),('Heading 3',10,DARK)]:
    s=styles[style_name]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(10); s.paragraph_format.space_after=Pt(5)
for name in ['List Bullet','List Bullet 2']:
    styles[name].font.name='Arial'; styles[name].font.size=Pt(9.2); styles[name].paragraph_format.space_after=Pt(2)

# Metadata
cp=doc.core_properties
cp.title='Anadolu Parsı Aile Yaşam Merkezi - Master Proje Dokümantasyonu - Bronze 04.08.2026.28'
cp.subject='Aktif kapsam, güvenlik, mimari, doğrulama ve teslim durumu'
cp.author='Panthera pardus tulliana'
cp.last_modified_by='Panthera pardus tulliana'
cp.keywords='Anadolu Parsı, Bronze, aile yaşam merkezi, proje dokümantasyonu'
cp.comments='Doğal kişi metadata içermez.'

# Header/footer
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=header.add_run('ANADOLU PARSI  |  BRONZE 04.08.2026.28'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=footer.add_run('Panthera pardus tulliana - Aktif proje belgesi'); r.font.name='Arial'; r.font.size=Pt(7.5); r.font.color.rgb=RGBColor.from_string(GRAY)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(55)
r=p.add_run('ANADOLU PARSI'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(17); r.font.color.rgb=RGBColor.from_string(GOLD)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Aile Yaşam Merkezi'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(29); r.font.color.rgb=RGBColor.from_string(TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16)
r=p.add_run('Master Proje Dokümantasyonu'); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor.from_string(DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Bronze 04.08.2026.28'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(GOLD)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(28)
r=p.add_run('Aktif yönetişim, kabul edilmiş kapsam, platform mimarisi, güvenlik politikası, uygulanan kod ve doğrulama durumu'); r.font.size=Pt(10.5)

rows=[['Alan','Değer'],['Üst marka','Panthera pardus tulliana'],['Uygulama','Anadolu Parsı Aile Yaşam Merkezi'],['Görünür sürüm',release['current']['visibleRelease']],['İç releaseId',release['current']['releaseId']],['Ana kaynak',release['current']['parentRelease']+' kaynak zincirinden devam'],['Kural sicili',f"{rules['ruleCount']} kural - {rules['activeRuleCount']} aktif - {rules['supersededRuleCount']} değiştirilmiş"],['Kabul edilmiş gereksinim',str(len(reqs))],['Belge statüsü','AKTİF - Bronze geliştirme kaydı; Silver/Gold kanıtı değildir']]
table(rows,[2.0,4.7])
callout('Gerçeklik sınırı','Bu belge yalnız kaynakta bulunan kayıtları özetler. NOT_RUN, BLOCKED, PENDING ve UNAVAILABLE sonuçları PASS değildir.', 'FFF2E1')

doc.add_page_break()

# 1
heading('1. Belgenin Yetkisi ve Tek Aktif Kaynak',1)
doc.add_paragraph('Bu Master Proje Dokümantasyonu, Bronze 04.08.2026.28 kaynak ağacındaki kanonik sicillerden üretilmiştir. Tarihsel Build/RC/MVP belgeleri kanıt ve tarihçe olarak korunur; aktif kapsamı veya güncel sürüm adlandırmasını değiştiremez.')
rows=[['Öncelik','Yetkili kaynak','İşlev']]
for i,pth in enumerate(active_docs['authorityOrder'],1):
    rows.append([str(i),pth,'Aktif yönetişim/kapsam kaynağı'])
table(rows,[.55,4.25,1.9])
callout('Kalıcı kayıt yolu',governance['persistentLibraryReleasePath']+'\n/mnt/data yalnız geçici çalışma alanıdır; kalıcı teslim değildir.',LIGHT)

heading('1.1 Kural ve karar kilidi',2)
bullet(f"Kanonik kural sicili: {rules['id']} - SHA-256 {governance['canonicalRulesSha256']}")
bullet(f"Toplam {rules['ruleCount']} kural: {rules['activeRuleCount']} aktif, {rules['supersededRuleCount']} açıkça değiştirilmiş.")
bullet(f"Kullanıcı karar sicili: {len(decisions.get('decisions',[]))} kayıt; DEC-123-DEC-128 bu sürümde kabul edildi.")
bullet('Preflight sonrasında kaynak parmak izi değişirse build/test/paketleme kapısı kapanır.')
bullet('Tamamlanmış sürümler değiştirilemez; yeni çalışma aylık sıradaki sürümde yapılır.')

heading('1.2 Sürüm sözleşmesi',2)
table([
 ['Kural','Değer'],
 ['Görünür biçim','Bronze/Silver/Gold gg.aa.yyyy.aylık-sıra'],
 ['Mevcut sürüm',release['current']['visibleRelease']],
 ['Aylık sıra',str(release['current']['monthlySequence'])],
 ['Ay değişimi','Sayaç yeni ayda sıfırlanır; aynı ay gerçek geçmiş sayısından devam eder'],
 ['Yasak görünür etiketler','RC, RC2, MVP ve küresel Build numarası'],
 ['Tarihsel belgeler','Adları değiştirilmeden HISTORY_ONLY olarak korunur']
],[2.0,4.7])

# 2
heading('2. Kabul Edilmiş Ürün Kapsamı',1)
doc.add_paragraph('Kullanıcı tarafından bu konuşmaya kadar önerilen bütün yetenekler bağlayıcı Bronze kapsamına alınmıştır. Kapsam sicili her öğeyi ayrı kimlik, öncelik, gerçek durum, kabul ölçütü ve kanıt zinciriyle tutar.')
table([
 ['Durum','Adet','Anlam'],
 ['COMPLETE',status_counts['COMPLETE'],'Karar-kod-test-belge-kanıt zinciri tamamlanmış'],
 ['FOUNDATION_STARTED',status_counts['FOUNDATION_STARTED'],'Temel kod var; uçtan uca özellik tamamlanmadı'],
 ['PARTIAL',status_counts['PARTIAL'],'Bir veya daha fazla zorunlu katman eksik'],
 ['NOT_IMPLEMENTED',status_counts['NOT_IMPLEMENTED'],'Kodlanmadı']
],[2.0,1.0,3.7])
table([
 ['Öncelik','Adet'],
 ['P0',priority_counts['P0']],['P1',priority_counts['P1']],['P2',priority_counts['P2']]
],[2.0,1.1])

heading('2.1 Ana yetenek aileleri',2)
for text in [
'Aile, hane, aile dalı, üyelik, soy ağacı, zaman tüneli, önemli günler ve aile hafızası.',
'Arşiv, OCR, tam metin arama, şifreli dosya paylaşımı, içerik bütünlüğü ve türetilmiş veri yönetimi.',
'Finans, banka hesabı, IBAN, kart, kredi, varlık, bütçe, hedef, nakit akışı ve açık bankacılık adapterleri.',
'Sağlık, bakım, ileri yaş, ilaç, randevu, çocuk ve eğitim koordinasyonu.',
'Hane operasyonu, afet/acil durum, ev-araç-eşya-evcil hayvan, seyahat ve akıllı ev/enerji.',
'Mesajlaşma, çevrimiçi durum, görüntülü görüşme, aile toplantısı, kayıt/rıza, canlı altyazı ve dil çevirisi.',
'AI asistanı, aile hafızası stüdyosu, dijital miras, zaman kapsülü ve gizlilik/veri sahipliği merkezi.',
'Windows Core Service, çoklu Windows node, quorum/witness/failover, Apple companion API ve imzalı eklenti platformu.'
]: bullet(text)

heading('2.2 En büyük kapsam alanları',2)
rows=[['Alan','Gereksinim adedi']]+[[a,n] for a,n in area_counts.most_common(20)]
table(rows,[4.7,1.5])

# 3
heading('3. Platform Mimarisi',1)
heading('3.1 Hedef çalışma modeli',2)
for text in [
'Windows Core Service, kullanıcı arayüzünden bağımsız ve 7/24 çalışan otoritatif çekirdektir.',
'Windows Desktop yalnız yerel Core Service istemcisidir; UI kapanınca sunucu görevleri durmaz.',
'macOS, iPhone, iPad, Apple Watch ve Vision Pro ilk aşamada yetki filtreli companion istemcilerdir.',
'Her Windows node yalnız kendi yerel şifreli SQLite projectionını açar; canlı SQLite dosyası ağdan paylaşılmaz.',
'Otomatik failover için üç oy gerekir: üç tam node veya iki tam node + witness.',
'Quorum kaybında yazılar fail-closed durur; son doğrulanmış veri read-only gösterilebilir.',
'Replika yedek değildir; yerel disk, haricî disk ve OneDrive hedefleri bağımsız kalır.'
]: bullet(text)

heading('3.2 API ve iletişim katmanları',2)
table([
 ['Katman','Amaç','Güvenlik sınırı'],
 ['Local Administration','Desktop, installer ve bakım araçları','Yerel named pipe/loopback; kimlik doğrulamalı, sürümlü, boyut sınırlı'],
 ['Companion Client API','Apple ve diğer istemciler','HTTPS + mTLS + kullanıcı/cihaz/amaç politikası'],
 ['Cluster Replication','Windows node log/snapshot/file sync','mTLS gRPC/protobuf; yalnız cluster kimlikleri'],
 ['Control Plane','Rendezvous, iptal, APNs wake, witness','Minimum metadata; aile içeriğini çözemez'],
 ['Media Plane','WebRTC/SFU/TURN','SFrame E2EE; relay/SFU içerik çözemez']
],[1.55,2.25,2.9])

heading('3.3 Bu sürümde gerçek kodlanan platform temeli',2)
work_tr = {
'200-rule canonical registry and 8-decision ledger':'200 maddelik kanonik kural sicili ve 8 karar kayıtlı kullanıcı karar defteri',
'Fail-closed governed preflight/postflight and source fingerprint':'Fail-closed yönetişim preflight/postflight kapıları ve kaynak parmak izi',
'Exhaustive artifact/document inventory':'Eksiksiz kaynak, kanıt ve belge envanteri',
'Secure authenticated local Core Service administration protocol':'Güvenli ve kimlik doğrulamalı yerel Core Service yönetim protokolü',
'Signed Platform Policy decision receipts across local transport':'Yerel taşıma üzerinden imzalı Platform Policy karar makbuzları',
'Desktop local Core Service adapter with injected authority':'Dışarıdan güvenli yetki bilgisi alan Desktop yerel Core Service adaptörü'
}
for text in delivery['workCompleted']:
    bullet(work_tr.get(text,text))
callout('Henüz tamamlanmayan bağlantı','Desktop adapter kaynakta vardır fakat production başlangıç/IPC/UI composition zincirine henüz bağlanmamıştır. Windows Service Host, cihaz bağlı secret provisioning, consensus ve Apple istemciler de tamamlanmamıştır.','FFF2E1')

# 4
heading('4. Güvenlik, Gizlilik ve Politika Çekirdeği',1)
for text in [
'Her istek kullanıcı, kişi, cihaz, uygulama, servis, aile/hane/dal, veri sahibi, işlem, amaç, süre, rıza ve hassasiyet bağlamında değerlendirilir.',
'Açık ret, rol veya varsayılan iznin üzerindedir. Aile yöneticiliği başkasının özel sağlık, finans, konum, belge veya mesaj içeriğine otomatik erişim sağlamaz.',
'Politika kararı bulunmayan hassas işlem fail-closed reddedilir.',
'Her uygulama ve worker aynı Platform Policy Kernel kurallarını devralır.',
'OCR, AI, çeviri, transkript, özet, embedding, thumbnail ve cache kaynak verinin hassasiyetini ve saklama politikasını miras alır.',
'Dış AI/OCR/çeviri sağlayıcısına aktarım varsayılan kapalıdır; içerik önizlemesi ve ayrı onay gerekir.',
'Ses/görüntü/toplantı kaydı varsayılan kapalıdır; tüm katılımcılar bilgilendirilip açık onay vermeden başlayamaz.',
'Yeni kodda doğrudan rol kontrolü yasaktır. Kaynakta kalan 35 eski kontrol teknik borç baselineıdır ve artamaz.',
'Kalıcı kullanıcı verisi şifrelenir; kimlik doğrulama öncesi aile verisi açılmaz.',
'Loglarda hassas içerik, parola, token, OTP, tam kart numarası, CVV/CVC veya PIN tutulmaz.'
]: bullet(text)

heading('4.1 Core Service yerel yönetim güvenliği',2)
table([
 ['Kontrol','Uygulanan davranış'],
 ['Protokol','Sürümlü contract, method allowlist ve typed response'],
 ['Kimlik doğrulama','Enjekte edilen secret; SHA-256 digest ve constant-time karşılaştırma'],
 ['Mesaj sınırı','64 KiB maksimum, tek istek/bağlantı ve timeout'],
 ['Yetkilendirme','Platform Policy Kernel kararı + imzalı policy receipt'],
 ['Korelasyon','Request/correlation ID'],
 ['Quorum/safe mode','Yazı istekleri reddedilir'],
 ['Secret kullanımı','Environment/default credential yok; loglama yok']
],[2.0,4.7])

# 5
heading('5. Kullanıcı Deneyimi ve Erişilebilirlik',1)
for text in [
'Varsayılan gövde metni en az 16 px hedefler; kritik küçük yazılar kaldırılır.',
'Windows metin ölçeği %100-%225, uygulama ölçeği ve %100-%400 DPI/display ölçeği desteklenir.',
'1280x720’den 4K’ya, küçük pencereye ve çoklu monitöre reflow gerekir.',
'Klavye ile tam kullanım, mantıksal sekme sırası, görünür focus ve ekran okuyucu semantiği zorunludur.',
'Windows Narrator, Magnifier, high contrast/forced colors, color blindness ve reduced motion test edilir.',
'Renk tek başına anlam taşımaz; etkileşim hedefleri en az 44 px olur.',
'Genç, standart, ileri yaş, az gören ve bakım veren görsel profilleri yetkiden ayrı tutulur.',
'Taslak, otomatik kaydetme, geri alma, hata, boş, yükleniyor, çevrimdışı ve yeniden deneme durumları tamamlanır.',
'İletişimde canlı altyazı, gerçek zamanlı metin, işaret dili konuşmacısını sabitleme ve düşük bant genişliği modu bulunur.'
]: bullet(text)

# 6
heading('6. Yönetişim Kapıları ve Aşılamazlık Modeli',1)
table([
 ['Kapı','Amaç','Mevcut sonuç'],
 ['Canonical Rule Registry','200 kuralın tekil/aktif/superseded bütünlüğü','PASS'],
 ['User Decision Ledger','Kullanıcı kararlarının kural ve belge bağlantısı','PASS'],
 ['Active Release Contract','Sürüm, workspace ve görünür adlandırma','PASS'],
 ['Governed Preflight','Kural hash, karar onayı ve kaynak parmak izi','PASS'],
 ['Preflight Tamper Runtime','Preflight sonrası kaynak değişirse kapının kapanması','PASS'],
 ['Feature Reality','Eksik özelliğin tamamlandı gösterilmemesi','PASS - Silver BLOCKED'],
 ['Platform Policy','Yeni bypass ve politika runtime doğrulaması','PASS'],
 ['Core Service Contract/Runtime','Yerel yönetim protokolü ve çalışma davranışı','PASS'],
 ['Document/Artifact Index','Bütün dosya ve belgelerin sınıflandırılmış dizini','PASS'],
 ['Clean npm ci','Temiz bağımlılık kurulumu','BLOCKED - OPEN-002'],
 ['Full TypeScript / full tests / Electron build','Tam kaynak doğrulaması','NOT_RUN'],
 ['Windows installer/launch/UAT','Gerçek Windows kabulü','NOT_RUN']
],[1.7,3.2,1.8])

heading('6.1 Preflight sonrası kaynak kilidi',2)
doc.add_paragraph('Preflight, yönetilen kaynak/config/aktif belge dosyalarının içerik tabanlı parmak izini üretir. Build, test ve paketleme komutları bu parmak izinin güncel olduğunu doğrular. Tek bir kaynak değişikliği bile preflightı geçersiz kılar ve işlem durur. Tamper testi geçici bir dosya ekleyerek kapının gerçekten kapandığını, dosya kaldırıldığında yeniden açıldığını doğrulamıştır.')

# 7
heading('7. Gerçek İlerleme ve Tahmini Takvim',1)
table([
 ['Gösterge','Değer'],
 ['Bronze tamamlanma',f"%{progress['codingCompletionPercent']:.1f}"],
 ['Bronze kalan',f"%{progress['codingRemainingPercent']:.1f}"],
 ['Makul aralık',f"%{progress['credibleRangePercent'][0]}-%{progress['credibleRangePercent'][1]}"],
 ['Tahmin güveni',{'LOW':'DÜŞÜK','MEDIUM':'ORTA','HIGH':'YÜKSEK'}.get(progress['confidence'],progress['confidence'])],
 ['Bronze kalan odaklı iş',f"{progress['remainingFocusedEngineeringWorkdays']['bronze'][0]}-{progress['remainingFocusedEngineeringWorkdays']['bronze'][1]} iş günü"],
 ['Tahmini Bronze bitişi',f"{progress['estimatedBronzeCompletion']['from']} - {progress['estimatedBronzeCompletion']['to']}"],
 ['Silver test/düzeltme',f"{progress['remainingFocusedEngineeringWorkdays']['silver'][0]}-{progress['remainingFocusedEngineeringWorkdays']['silver'][1]} iş günü"],
 ['Tahmini Silver aralığı',f"{progress['estimatedSilverTransitionAndCompletion']['from']} - {progress['estimatedSilverTransitionAndCompletion']['to']}"],
 ['Gold kapanışı',f"{progress['remainingFocusedEngineeringWorkdays']['gold'][0]}-{progress['remainingFocusedEngineeringWorkdays']['gold'][1]} iş günü"],
 ['Tahmini Gold hazırlığı',f"{progress['estimatedGoldReadiness']['from']} - {progress['estimatedGoldReadiness']['to']}"],
 ['Silver durumu','YASAK / HAZIR DEĞİL'],
 ['Gold durumu','YASAK / HAZIR DEĞİL']
],[2.7,4.0])
callout('Tahmin sınırı','Bu tarihler tek odaklı geliştirme akışı, dış kimlik/entegrasyon/test hesaplarının zamanında sağlanması ve yeni büyük kapsam eklenmemesi varsayımına dayanır. Bunlar taahhüt değil, düşük güvenli mühendislik tahminidir.','FFF2E1')

heading('7.1 Sohbet alanı ve devam promptu',2)
table([
 ['Alan','Durum'],
 ['Gerçek sohbet ölçümü','%9,4 kullanılmış / %90,6 kalan - platform actual ölçümü'],
 ['Tahmin üretme','YASAK - yalnız actual ölçüm kullanıldı'],
 ['Devam promptu','GEREKMİYOR - %90 kullanım hard-stop eşiği oluşmadı'],
 ['Hard-stop oluşursa','Kanonik kaynak, Library yolu, son sürüm, SHA, açık işler ve sıradaki görevle otomatik devir promptu üretilir']
],[2.2,4.5])

# 8
heading('8. Belge, Kaynak ve Kanıt Envanteri',1)
doc.add_paragraph('Kullanıcıya her dosyayı sohbet içinde binlerce satır hâlinde tekrarlamak yerine, makine okunur ve insan okunur eksiksiz dizinler üretilir. Dizin kapısı, kaynak ağacındaki her dosyanın listede yer aldığını doğrular.')
table([
 ['Envanter','Dosya'],
 ['Bütün dosya/artifact dizini','artifacts/manifests/PROJECT_ARTIFACT_INDEX.json / .csv / .md'],
 ['Bütün belge dizini','artifacts/manifests/ALL_DOCUMENTS_INDEX.json / .csv / .md'],
 ['Kullanıcıya yönelik belge dizini','docs/current/08_TUM_BELGELER_DIZINI.md'],
 ['Aktif belge önceliği','config/active-document-set.json'],
 ['Kanonik kural sicili','config/canonical-rule-registry.json'],
 ['Kabul edilmiş kapsam','config/accepted-scope-registry.json'],
 ['Kullanıcı kararları','config/user-decision-ledger.json']
],[2.25,4.45])

heading('8.1 Bu sürümde oluşturulan veya değiştirilen ana belgeler',2)
for pth in [
'config/canonical-rule-registry.json','config/project-constitution.json','config/active-governance-ledger.json','config/user-decision-ledger.json','config/accepted-scope-registry.json','config/release-ledger.json','config/delivery-report-contract.json','config/conversation-capacity-policy.json','config/persistent-artifact-policy.json','config/document-inventory-policy.json','docs/current/00_AKTIF_ANA_KAPSAM.md','docs/current/01_AKTIF_GUVENLIK_GIZLILIK_POLITIKASI.md','docs/current/02_AKTIF_PLATFORM_MIMARISI.md','docs/current/03_AKTIF_UX_ERISILEBILIRLIK_SOZLESMESI.md','docs/current/04_AKTIF_BRONZE_YOL_HARITASI.md','docs/current/05_KABUL_EDILEN_EK_AILE_YETENEKLERI.md','docs/current/06_KANONIK_KURAL_SICILI.md','docs/current/07_TESLIM_SOHBET_VE_KALICI_KAYIT_SOZLESMESI.md','docs/current/08_TUM_BELGELER_DIZINI.md','docs/current/09_KULLANICI_KARARLARI_KAYDI.md','artifacts/reports/PROJECT_PROGRESS_04.08.2026.28.json','artifacts/reports/DELIVERY_STATUS_04.08.2026.28.json']:
    bullet(pth)

# 9
heading('9. Açık Riskler ve Sonraki Resmî İş',1)
risk_tr = {
'35 legacy direct family_admin checks remain and are blocked from increasing':'Kaynakta 35 eski doğrudan family_admin kontrolü kalmıştır; sayılarının artması otomatik kapıyla engellenir.',
'Desktop adapter is not yet wired into production IPC/UI composition':'Desktop adaptörü production başlangıç, IPC ve UI composition zincirine henüz bağlanmamıştır.',
'Windows Service Host, device-bound local admin secret provisioning, cluster consensus and Apple clients are not implemented':'Windows Service Host, cihaza bağlı yerel yönetim secret provisioning, cluster consensus ve Apple istemcileri henüz kodlanmamıştır.',
'Full dependency/typecheck/build/Windows validation remains unavailable or not run':'Tam bağımlılık kurulumu, typecheck, build ve Windows doğrulaması kullanılamıyor veya çalıştırılmadı.',
'Persistent Library upload occurs only after deterministic source packaging':'Kalıcı Library yüklemesi yalnız deterministik kaynak paketlemesi tamamlandıktan sonra yapılır.'
}
for risk in delivery['openErrorsAndRisks']:
    bullet(risk_tr.get(risk,risk))
callout('Sonraki tek resmî kodlama işi','Desktop başlangıç ve sistem sağlığı IPC akışını kimlik doğrulamalı yerel Core Service istemcisine bağla; ardından ilk hassas veri use-case’ini imzalı policy receipt zorlamasının arkasına taşı.',LIGHT)

heading('9.1 Yayın geçiş kuralı',2)
for text in [
'Bronze, kabul edilmiş bütün özelliklerin kodlandığı geliştirme kanalıdır.',
'Bronze kapsamındaki P0/P1 eksikleri kapanmadan Silver açılamaz.',
'Silver yalnız tam test, UAT, güvenlik, performans, erişilebilirlik ve Windows kurulum doğrulaması ile hata düzeltme kanalıdır.',
'Silver sırasında yeni özellik boşluğu bulunursa çalışma Bronze’a döner.',
'Gold ancak bütün Silver kapıları PASS, üretim imzalama/operasyon tamam ve ürün sahibi onayı mevcutsa açılır.'
]: bullet(text)

# final page
heading('10. Zorunlu Bitiş Durumu',1)
table([
 ['Alan','Durum'],
 ['Kanal ve sürüm',release['current']['visibleRelease']],
 ['Tamamlanan gereksinimler',', '.join(delivery['completedRequirementIds'])],
 ['Tamamlanan kararlar',', '.join(delivery['completedDecisionIds'])],
 ['Bronze ilerleme',f"%{progress['codingCompletionPercent']:.1f} tamamlandı - %{progress['codingRemainingPercent']:.1f} kaldı"],
 ['Silver','YASAK / HAZIR DEĞİL'],
 ['Gold','YASAK / HAZIR DEĞİL'],
 ['Sohbet doluluğu','%9,4 kullanılmış / %90,6 kalan - actual ölçüm'],
 ['Kalıcı Library hedefi',governance['persistentLibraryReleasePath']],
 ['Kaynak ZIP/SHA','Paketleme sonrasında detached teslim tasdikinde kesinleşir'],
 ['Sonraki görev','Desktop başlangıç/sistem sağlığı IPC bağlantısı ve ilk hassas use-case için policy receipt zorlaması']
],[2.0,4.7])
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18)
r=p.add_run(delivery['mandatoryTruthSentence']); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(RED)

OUT.parent.mkdir(parents=True,exist_ok=True)
doc.save(OUT)
print(OUT)
