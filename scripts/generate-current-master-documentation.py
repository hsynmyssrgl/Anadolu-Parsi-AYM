from __future__ import annotations

import hashlib
import html
import json
import re
import subprocess
from collections import Counter
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    CondPageBreak,
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
AS_OF = "27.08.2026"
VERSION = "GUNCEL-2026-08-27-V5"
SOURCE = ROOT / "docs/current/11_GUNCEL_KARAR_KURAL_IS_AKISI_SICILI.md"
DOCX_OUT = ROOT / "docs/current/MASTER_PROJE_DOKUMANTASYONU_GUNCEL_24.08.2026_V5.docx"
PDF_OUT = ROOT / "docs/current/MASTER_PROJE_DOKUMANTASYONU_GUNCEL_24.08.2026_V5.pdf"
LOGO = ROOT / "apps/desktop/src/renderer/assets/brand-mark.png"

BRONZE = "7C4D20"
BRONZE_LIGHT = "F5EEE7"
GOLD = "D5A526"
GREEN = "467259"
INK = "333537"
MUTED = "666B69"
RED = "A33A3A"
BLUE_GRAY = "E8EEF5"


def read_json(path: str):
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


def read_text(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def first_heading(path: Path) -> str:
    for line in path.read_text(encoding="utf-8-sig", errors="replace").splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    return path.stem


def status_from_markdown(path: Path) -> str:
    for line in path.read_text(encoding="utf-8-sig", errors="replace").splitlines():
        normalized = line.replace("**", "").strip()
        match = re.match(r"^(?:[-•]\s*)?(?:Durum|Status)\s*:\s*(.+)$", normalized, flags=re.I)
        if not match:
            continue
        value = match.group(1).strip().strip("`").strip()
        if ". " in value:
            value = value.split(". ", 1)[0].strip()
        value = value.replace("`", "").strip()
        return visible_status(value or "KAYITLI")
    return "KAYITLI"


def visible_status(value) -> str:
    return str(value or "-").replace("_", " ")


def decision_number(path: Path) -> int:
    match = re.match(r"DEC-(\d+)", path.name)
    return int(match.group(1)) if match else 9999


def adr_number(path: Path) -> int:
    match = re.match(r"ADR-(\d+)", path.name)
    return int(match.group(1)) if match else 9999


def git_head() -> str:
    return subprocess.check_output(["git", "rev-parse", "HEAD"], cwd=ROOT, text=True).strip()


def get_requirement_counts(registry: dict) -> dict[str, int]:
    return dict(Counter(item.get("status", "UNKNOWN") for item in registry["requirements"]))


def scan_existing_word_pdf() -> list[dict]:
    rows = []
    excluded = {DOCX_OUT.name, PDF_OUT.name}
    for path in sorted((ROOT / "docs/current").iterdir()):
        if path.name in excluded or path.suffix.lower() not in {".docx", ".pdf"}:
            continue
        row = {"name": path.name, "type": path.suffix.lower()[1:], "bytes": path.stat().st_size, "readable": True}
        try:
            if path.suffix.lower() == ".pdf":
                reader = PdfReader(str(path))
                row["pages"] = len(reader.pages)
                row["characters"] = sum(len(page.extract_text() or "") for page in reader.pages)
            else:
                document = Document(path)
                row["paragraphs"] = len(document.paragraphs)
                row["tables"] = len(document.tables)
                with ZipFile(path) as archive:
                    app_xml = archive.read("docProps/app.xml").decode("utf-8", errors="ignore")
                    pages = re.search(r"<Pages>(\d+)</Pages>", app_xml)
                    if pages:
                        row["pages"] = int(pages.group(1))
        except Exception as error:  # pragma: no cover - audit-only branch
            row["readable"] = False
            row["error"] = str(error)
        rows.append(row)
    return rows


ledger = read_json("config/active-governance-ledger.json")
rules = read_json("config/canonical-rule-registry.json")
constitution = read_json("config/project-constitution.json")
requirements = read_json("config/accepted-scope-registry.json")
roadmap = read_json("config/remaining-scope-package-roadmap.json")
visual = read_json("config/ui-visual-reference-manifest.json")
closure = read_json("config/34-l-bronze-final-drift-deterministic-delivery-closure-scope.json")
windows_scope = read_json("config/34-k-windows-resilience-universal-ux-scope.json")
document_index = read_json("artifacts/manifests/ALL_DOCUMENTS_INDEX.json")
full_document_audit = read_json("artifacts/manifests/ALL_PROJECT_DOCUMENT_FORMAT_AUDIT.json")
user_decisions = read_json("config/user-decision-ledger.json")
requirement_counts = get_requirement_counts(requirements)
decision_counts = dict(Counter(item.get("status", "UNKNOWN") for item in user_decisions["decisions"]))
decision_files = sorted((ROOT / "docs/decisions").glob("DEC-*.md"), key=decision_number)
adr_files = sorted((ROOT / "docs/adr").glob("ADR-*.md"), key=adr_number)
threat_files = sorted((ROOT / "docs/security").glob("*.md"))
word_pdf_scan = scan_existing_word_pdf()
readable_count = sum(1 for row in word_pdf_scan if row["readable"])
historical_pairs = len(word_pdf_scan) // 2


WORKFLOWS = [
    ("Karar ve kural değişikliği", "Açık kullanıcı kararı → aynı değişiklikte DEC kaydı + makine defteri + etkilenen aktif belgeler + iş listesi açık/kapalı/neden alanları → etki analizi → gerekiyorsa kanonik kural/scope ve kod → doğrulama → indeks/master DOCX/PDF. Eşzamanlılık eksikse karar veya iş tamamlandı sayılamaz; sessiz istisna ve waiver yoktur."),
    ("Paket yürütme", "Aynı anda tek paket/adım IN_PROGRESS olabilir. Paket, yerel PASS üretse bile dış ve manuel kabul kanıtı eksikse açık kalır. Sonraki paketin başlaması, öncekinin kalıcı receipt ve kabul koşullarına bağlıdır."),
    ("Yetkilendirilmiş işlem", "Oturum kimliği + hesap + kişi + aile + cihaz + güvenlik epoch + amaç + kaynak + hassasiyet merkezi PEP tarafından değerlendirilir; receipt/UoW/optimistic revision/audit/outbox aynı yönetilen işlem sınırında korunur."),
    ("Veri yaşam döngüsü", "Kaynak sınıflandırma → sahiplik ve saklama → erişim/işleme gözlemi → türetilmiş veri mirası → geri alınabilir silme → kaynak silme yayılımı → içeriksiz tombstone → yedek/haricî kopya için açık risk. Fiziksel güvenli silme garantisi verilmez."),
    ("Kimlik ve cihaz", "Passkey challenge ve doğrulama → cihaz/epoch bağlama → kayıp anahtar kurtarma → yerel oturum iptali. Federated kimlik yalnız yapılandırılmış ve güvenilen sağlayıcı/JWKS/ağ zinciriyle görünür; gerçek sağlayıcı testi yoksa PASS yoktur."),
    ("Yerel OCR", "Archive kaynağı PEP + ayrı hassas işleme rızası → main-only byte okuma → bounded child process → malware kapısı → sealed sonuç kasası → düzeltme/rerun lineage → source deletion propagation. PDF/malware/low privilege eksikleri fail-closed kalır."),
    ("İletişim", "Kimlik/politika → MLS/E2EE oturum → içerikten ayrı audit → mesaj/dosya yaşam döngüsü → çağrı preflight → açık kayıt rızası → medya retention → çeviri/altyazı. Gerçek provider/cihaz/ağ UAT olmadan üretim iddiası yoktur."),
    ("Dağıtık çalışma", "Windows tek-yazar yerel veri otoritesi → mutation log → node kimliği/mTLS → quorum/witness → snapshot/failover → istemci API/cache → operasyon ve DR. Yerel simülasyon gerçek çoklu node kanıtı değildir."),
    ("Windows teslim", "Governed preflight → typecheck/test/build → Electron fuse/ASAR doğrulama → imzalama/provenance → installer → kurulu uygulama açılışı → update/rollback → uninstall/residue → postflight. Sertifika veya lifecycle UAT eksikse dağıtım hazır sayılmaz."),
    ("Belge güncelleme", "Canlı kaynak ve JSON sicilleri taranır → açık/kapalı/neden matrisi güncellenir → Markdown tek kaynak oluşturulur → DOCX/PDF aynı veriden üretilir → bütün sayfalar render edilir → hash manifesti yenilenir → tarihsel belgeler korunur."),
]


INFRASTRUCTURE = [
    ("Platform Policy Kernel", "Fail-closed merkezi karar; doğrudan rol kontrolü yeni kodda yasak; receipt/UoW/audit/outbox bağları."),
    ("Windows Desktop", "Electron main güvenlik otoritesi, preload exact bridge, renderer yetkisiz sunum katmanı, korumalı yerel secret/artifact sınırları."),
    ("Core Service", "Ayrı Electron utility companion süreci, sürümlü typed API, CurrentUser DPAPI korumalı provisioning, her açılışta yeni local named-pipe/token ve içeriksiz lifecycle protokolü; aile verisi/SQLite sahipliği Desktop'ta kalır, Windows Service/cluster kabulü ayrıca gerekir."),
    ("Veri katmanı", "Şifreli yerel SQLite, migration/trigger/optimistic revision/idempotency, append-only mutation/audit, source-deletion lineage."),
    ("Kimlik", "Passkey, bounded challenge, lost-key recovery, trusted OIDC+PKCE/JWKS/vault, temporary credential, companion snapshot temelleri."),
    ("OCR", "PEP-authorized archive read, bounded child process, malware fail-closed, sealed result vault, correction lineage ve purge."),
    ("İletişim", "MLS/E2EE metadata/payload sınırları, messaging, presence, calling, recording consent, translation, meetings, file sharing, archive audit."),
    ("Dağıtık platform", "Mutation/consensus/tenancy, distributed clients/operations/DR ve Windows resilience için local hardened foundation."),
    ("Görsel sistem", "Onaylı sıcak-nötr açık tema, exact release colors, 512 px sıcak-bronz logo, merkezi typography/accessibility tokenları."),
    ("Belge yönetişimi", "DEC-251 eşzamanlı karar-belge-iş listesi kapısı; tüm belge türü hash/okunabilirlik denetimi; tarihsel kayıtların değişmezliği."),
    ("Marka ve kurumsallaşma", "DEC-261 ile ana marka ParsYuva, görünür ürün ParsYuva Aile Yaşam Merkezi; AYM yalnız tarihsel ve zorunlu teknik uyumluluk kimliklerinde korunur. Şirket, marka, alan adı, hukuk/vergi/gizlilik ve mağaza kanıtları NOT_RUN iken tamamlanmış gösterilmez."),
]


EXTERNAL_DEPENDENCIES = [
    ("Üretim sertifikası ve provenance", "Authenticode/kod imzalama sertifikası, güvenilir zaman damgası ve üretim provenance kanıtı yoksa installer yayıma hazır sayılamaz."),
    ("Gerçek cihaz ve sağlayıcı UAT", "Passkey/authenticator, Windows cihazları, kamera/mikrofon, Matter, OCR/AI/çeviri, OIDC, WebRTC/SFU/TURN, bulut/kurum bağlantıları ve Apple istemcileri gerçek ortamda sınanmalıdır."),
    ("Gerçek dağıtık sistem", "En az gerçek çoklu node, quorum/witness, mTLS, ağ bölünmesi, failover, snapshot ve felaket kurtarma provası gerekir."),
    ("Uzun süreli işletim", "168 saat soak, yeniden başlatma, güç kesintisi, disk doluluğu, saat değişimi, uyku/uyanma ve güncelleme/rollback kanıtları gerekir."),
    ("Hukuk ve gizlilik", "Saklama/imha süreleri, çocuk/sağlık/iletişim kaydı, delil niteliği, sağlayıcı sözleşmeleri ve ülke bazlı yükümlülükler uzman incelemesi olmadan ürün gerçeği olarak sunulamaz."),
    ("İnsan ve erişilebilirlik UAT", "Narrator, büyütme, klavye, contrast, metin taşması, çocuk/yetişkin/ileri yaş/bakım veren profilleri ve gerçek kritik akışlar insanlarla doğrulanmalıdır."),
]


DRIFT_FIXES = [
    "Aktif kapsam toplamı 344 → 358 olarak düzeltildi ve durum dağılımı eklendi.",
    "Kanonik kural sicili V10/214/194 ve güncel SHA ile ParsYuva marka, kurumsallaşma, platform ve belge sınıflandırma kurallarına yükseltildi.",
    "Kullanıcı karar defteri 83 kayda yükseltildi; DEC-254 marka uyumluluğu ve kurumsallaşma kararını bağladı.",
    "Platform mimarisindeki 'OCR/iletişim başlamadı' anlatımı yerel bileşim var fakat kabul dış kanıta bağlı şeklinde düzeltildi.",
    "Yol haritasına her açık paket için yerel uygulama durumu, açık kalma nedeni, eksik kanıt ve requirement PASS alanları eklendi.",
    "Görsel sözleşme 17 px body, exact Bronze/Silver/Gold tokenları ve onaylı 512 px sıcak-bronz logo SHA'sıyla hizalandı.",
    "Geçici .tmp/.tmp-runtime-dist içeriğinin aktif belge envanterine girmesi engellendi.",
    "Build209–228 ve eski Bronze master çiftleri tarihsel olarak korundu; yeni sürüm ayrı adla oluşturuldu.",
    "Word/PDF dışındaki RTF, Markdown, JSON/YAML, TXT, CSV ve HTML kayıtları da kök klasör düzeyinde hash ve okunabilirlik denetimine alındı; Excel/PowerPoint bulunmadığı açıkça kaydedildi.",
    "Her yeni kararın DEC, makine defteri, etkilenen belgeler ve açık/kapalı iş gerekçeleriyle aynı değişiklikte güncellenmesi fail-closed kurala bağlandı.",
    "Bu kapsamlı tarama tarihsel kayıtların son içerik temelidir; DEC-252 gereği gelecekte eski build/arşiv/checkpoint içeriği yeniden denetlenmeyecek, yalnız değişmez HISTORICAL kayıt olarak korunacaktır.",
    "Core Service companion ASAR paketine bağlandı; güncel tam regresyonda 350/350 test dosyası ve 2187/2187 test, root typecheck/build ve aynı profilde iki ardışık normal win-unpacked açılışı PASS verdi. Production Authenticode sertifikası bulunmadığından signed installer/kurulu yaşam döngüsü açık bırakıldı.",
    "EK-001–EK-019 tarihsel karar tamponu DEC-260 ile ana sicillere bağlandı; daha yeni ParsYuva, dil ve kurulum kararları çatışmada üstün tutuldu.",
    "Kanonik kural sicili V16/228/207 durumuna yükseltildi; tam ParsYuva Aile Yaşam Merkezi adı, sürüm paleti, parola görünürlüğü, installer yaşam döngüsü, aylık build, deneme/Gold, kaldırma-sıfırlama, tepsi ve migration/rollback kararları fail-closed kapılara bağlandı.",
    "DEC-261 ile AYM kısaltması güncel kullanıcı yüzeylerinden kaldırıldı; yalnız tarihsel kayıtlar ve değiştirilemeyen teknik uyumluluk yolları güncel marka olmadığı açıkça belirtilerek korunur.",
    "DEC-262 ile Windows kurulum hedefi C:\\Program Files\\PPT\\ParsYuva, ana program ve kısayol adı ParsYuva, teslim adı ParsYuva-<Kanal>-GG.AA.YYYY.NN.exe olarak sabitlendi.",
    "DEC-271 ile güncel kanal program kökleri legacy dizinin dışındaki C:\\Program Files\\PPT\\ParsYuva-<Kanal> kardeş yollarına taşındı; AppData ParsYuva/<Kanal> ve diğer kanal yalıtımı korunurken otomatik legacy veri migration veya silme yasaklandı.",
    "DEC-272 ile sürüm tahsisi exact expected release ID alan açık tek mutasyon oldu; preview yazmaz, uyuşmazlık yazım ve temizliğe geçmeden durur, signed/local/dir yalnız önceden tahsisli exact current kimliğini tüketir.",
    "DEC-273 ile Windows installer teslimi metadata-only kanonik UAT110 gerçek N→N+1 ve same-version maintenance koruması ile source/package/expected release bağlı schema2 kurulu ön yüz UAT111 makbuzuna bağlandı.",
    "DEC-274/PR-239 ile Windows teslim zinciri canlı PR-235 readback, schema2 package provenance, Bronze 50 bootstrap veya normal Bronze 52+ exact sibling continuation modlu UAT110 V3, installer-experience V2, parent-run bağlı UAT111 V3 ve final V3 geri-okuma kapılarıyla adversarial olarak güçlendirildi; exact Bronze 51 recovery modu DEC-276/PR-241 ile ayrıca ayrıldı.",
    "PR-235 ile en küçük değişiklik dahi exact değişen yol, bağımlı kural/karar/belge/manifest/ratchet/test/UAT kayıtları ve aynı temiz committe hedefli-tam-bütünlük kanıtlarıyla fail-closed eşlemeye bağlandı.",
    "PR-235 BOOTSTRAP_ADOPTION diff tabanı sabit kalırken producer yalnız pointer sourceCommit kayıt commitinde external-pointer exact eşitliği ve baseCommit → pointer.sourceCommit → HEAD ancestry ile; normal PRE_MUTATION producer ise kendi baseline commitinde doğrulanır.",
    "PR-239 UAT111 kapsamı Git'te izlenen TypeScript kanonik rota otoritesi, tüm görünür ve uygun kontrollerin dinamik outcome matrisi, gerçek native CANCEL/ACCEPT ve exclusive reparse-korumalı kanıt köküyle güncellendi.",
    "PR-240 filtresiz tam regresyon guard hatası 4c6652e0 ile korunmuştur. PPK-022 masaüstü başlangıç zincirinde --no-write aktarımı çağrı, sarmalayıcı ve son makbuz üreticilerinde kapatılmış; kaynak regresyonu 1 dosya/6 test, çalışma 51/51, sözleşme 41/41 ve uçtan uca PPK-022 24/24 PASS verirken 1.571 doğrulama dosyasında sıfır değişiklik kanıtlanmıştır. Bundled render aracının yinelenen LibreOffice PATH hatası 6ec632c8 ile korunmuş, mutlak LibreOffice/Poppler yollarıyla ana belge 28/28 sayfa görsel QA PASS vermiştir.",
    "PR-240 hedefli test üreticisinin zorunlu açık test listesi verilmeden çağrılması 4c8b6b7d ile fail-closed korunmuştur. Etki değerlendirmesi ve analizi 577 değişen yol ile 94 test dosyasını PASS hesaplamıştır; yeni exact koşu bu 94 dosyayı analizden türetilen sırayla açıkça aktarır ve ret PASS olarak yorumlanmaz.",
    "PR-240 kanonik Git-index üretimi ile varsayılan canlı-ağaç doğrulama modu uyuşmazlığı 47f441e1 ile fail-closed korunmuştur. Eş --git-index --no-report doğrulaması 13.146 kontrol/4.407 dosya/2.143 belge ve kaynak bütünlüğü 4.868/4.868 dosya/4.869 SHA PASS vermiştir; yanlış-mod ret PASS değildir.",
    "PR-240 Bronze runtime önkoşul zinciri 703be65a, 0a118f5f ve 7fb288cd ile; 33-Y/33-Z/34-A alt süreç no-write sızıntısı ddb1abff ile fail-closed korunmuştur. Paket/core-service/desktop çıktılarından sonra 16/16 ek runtime PASS; ortak no-write aktarımından sonra 1 dosya/6 test, üç runtime ve üç byte-exact manifest SHA karşılaştırması PASS vermiştir. Yeni exact commit tam regresyonunun yerine geçmez.",
    "PR-240 güncel master DOCX ilk PNG renderında bundled Poppler yolunun çözülememesi dd675310 ile fail-closed korunmuştur. Exact Poppler/LibreOffice retry 28/28 sayfa üretmiş; 10–25 byte-identical, değişen 1–9 ve 26–28 sayfalar özgün çözünürlükte görsel QA PASS vermiştir.",
    "PR-240 release kaynak bütünlüğü doğrulayıcısının ana app çalışma ağacında çağrılması 99ad48dd ile fail-closed korunmuştur. Bu ret ürün kaynak bozulması değildir; exact Bronze çalışma ağacındaki retry PASS olmadan paket üretilemez.",
    "PR-240 final master DOCX 0669cb38 kaynak commitinden exact bundled LibreOffice/Poppler yollarıyla 28/28 sayfa render ve görsel QA PASS vermiştir; 2–6 byte-exact aynı, değişen 1 ve 7–28 özgün çözünürlükte kusursuzdur.",
    "PR-240 exact 80cf2a39 Bronze koşusunda hedefli 94 dosya/598 test ve filtresiz 398 dosya/2.469 test PASS sonrasında 33-R alt sürecinin migration manifestinde yalnız generatedAt değiştirmesi c7a3c130 ile fail-closed korunmuştur. Dört alt doğrulayıcıya no-write aktarımı eklendikten sonra odaklı 1 dosya/6 test ve gerçek 33-R matrisi 11/11, 8 dosya/30 test PASS; manifest SHA byte-exact değişmezdir. Yeni temiz exact commit tam regresyonunun yerine geçmez.",
    "PR-240 exact bfb6db9f Bronze koşusunda hedefli 94 dosya/598 test ve filtresiz 398 dosya/2.469 test PASS iken üreticinin direct-node çağrısı npm ortam bağını kaldırdığı için 171 ek komuttan 34-B/C/D/F FAIL olmuş ve 51316ac3 ile korunmuştur. Kanonik npm ortamında 34-B 13/13, 5 dosya/30 test PASS; final tur yalnız resmi npm scriptiyle yeniden çalıştırılır.",
    "PR-240 c02744cd exact Bronze resmi npm tam regresyonu 398 dosya/2.469 test ve 171/171 ek komut PASS vermiştir. Kaynak bütünlüğü 683 eksik Git-dışı manifest payloadı ve linked-worktree .git yönetim dosyasının yanlış kaynak sayılması nedeniyle 684 bulguyla FAIL olmuş ve 7d67fcff ile korunmuştur. Kanal kurulumu üç ayrı kanala güvenli yol/normal dosya/byte/SHA-256/atomik readback bağlı payload eşitlemesi yapar; kaynak toplayıcı .git dosya ve klasörünü dışlar. Odaklı 2 dosya/9 test PASS; yeni exact kapılar zorunludur.",
    "PR-240 96b9faac kaynaklı güncel ana DOCX ilk PNG renderında yanlış bundled PATH nedeniyle Poppler pdfinfo çözülememiş ve e5787764 ile fail-closed korunmuştur. Doğru native\\poppler\\Library\\bin ve LibreOffice yollarıyla retry 28/28 sayfa üretmiş; bütün sayfalar temas sayfalarında, metin yoğun 6–7 ve tablo yoğun 27–28 ayrıca özgün çözünürlükte taşma, örtüşme, kesilme, font ve bozuk karakter açısından görsel QA PASS vermiştir.",
    "PR-240 kapanış kayıtlarını içeren 7f866e69 kaynak commitinden final master DOCX doğru native Poppler/LibreOffice zinciriyle 29/29 sayfa görsel QA PASS vermiştir. Tüm sayfalar beş temas sayfasında; 6–7 ve yoğun kural/envanter tablolarını taşıyan 26–29 ayrıca özgün çözünürlükte kusursuzdur.",
    "PR-240 d421c299 exact Bronze turunda hedefli 95 dosya/600 test, filtresiz 399 dosya/2.471 test, 172/172 ek komut ve 4.869/4.869 kaynak bütünlüğü PASS sonrasında governed preflight; önce güncel retention/görünür sürüm makbuzu eksikliğini 607a9a53, ardından tamamlanmış çalışma adımlarının 1.428 kanıt yolundan 803 Git-dışı checkpoint payload dosyasının kanal eksikliğini 8b2b5ccc ile fail-closed korumuştur. Kanal kurulumu tamamlanmış localEvidence ve persistent receipt yollarını tracked/manifest dışlamalı kanonik yol, normal dosya, SHA-256, atomik yazım ve readback ile üç kanala eşitler. Seçici 809 yol, odaklı 1 dosya/9 test ve ticari temel 1.234 kontrol PASS; yeni exact preflight PASS olmadan paket yoktur.",
    "PR-240 0f0a4653 exact Bronze etki analizi 95 hedefli test dosyası hesaplamış; hedefli turda 600 test PASS iken operation-rule-check-policy current-mutation preflightStatus alanındaki tarihsel 607a9a53/8b2b5ccc FAIL metnini reddetmiş ve 50f4d9e5 ile fail-closed korunmuştur. Tarihsel retler QA alanlarında kalır; güncel mutasyon durumu NOT_RUN_CURRENT_MUTATION olur. Bronze/Silver/Gold hidrasyonu her kanalda 1.428/1.428 ve eksik 0 PASS; yeni exact testler zorunludur.",
    "PR-241/DEC-276 Bronze 50 immutable REJECTED_INVALID_PACKAGE geçmişini trusted runtime saymadan korur. b0615638, 3eec5426 ve 86602f7a tarihsel retlerdir. Exact etki değerlendirmesinin eşlenmemiş Windows paketleyici lockfile reddi 3e496f47, ters ruleIds assertionı 398de9c8 ile korunmuş; assertion düzeltmesi 3 dosya/14 test PASS vermiştir. Render çağrı retleri 75e4072c/0f98c7cc ve kural FAIL sonrasında başlayan generator çağrısı 9a370e51 ile korunur. Temiz 9f16699d kapanışı sonrasında pre-sync kanal reddi dbefb586 ile korunmuştur. Ana kaynak ve Bronze/Silver/Gold 63c55074 commitinde exact temiz eşitlik PASS; etki zinciri 105 değişen yol/19 hedef test hesaplamış ve hedefli 19 dosya/188 test PASS vermiştir. Filtresiz turda 399 dosya/2.480 test PASS iken yalnız PPK-015 üretim kaynak ratchet hash eşliği düşmüş, gerçek FAIL cc922201 ile korunmuştur. Canlı sınır 18 bölge/590 dosya/0 bulgu/2 adapter/3 amaçtır; ağ yetkisi değişmemiştir. Ratchet debfeecf460834f50cf328bff58b2c19ad94ef229610c4c829a35c4331ef235a özetine eşlenir. V5 DOCX sayfa 15 Durum token sarımı FAIL'i 787c5570 ile korunmuş; görünür karar durumları okunur boşluklarla sarmalandıktan sonra final belge 29/29 görsel QA PASS vermiştir. Yalnız Bronze 51 rejected-parent provenance bundle'ını history-only lineage olarak kullanıp temiz recovery fresh-install ve ayrı same-version maintenance yolunu çalıştırabilir; bütün exact test, kaynak bütünlüğü, preflight/postflight, paket ve kurulu uygulama UAT kapıları zorunludur.",
    "V5 belge becerisi başlangıcında tek format bekleyen yardımcıya birleşik docx,pdf değeri verilmesi içerik üretmeden reddedilmiş ve df92cdba ile korunmuştur. Bu ürün veya belge içerik kusuru ve PASS değildir; desteklenen tek docx retry PASS olmuş, üretim ancak bundan sonra başlatılmıştır. Kapanış renderındaki sayfa 15 Durum token-ortası bölünme 787c5570 ile korunmuş; üretici görünür durum değerlerindeki alt çizgileri izinli boşluklara dönüştürdükten sonra final DOCX 29/29 görsel QA PASS vermiştir.",
    "6f9139e1 exact Bronze zincirinde assessment 107 değişen yol/21 hedef test; hedefli 21 dosya/211 test ve filtresiz 399 dosya/2.481 test PASS vermiştir. Source-integrity governed preflight sonrasında yenilenen yedi indeks/dizin dosyası ile ticari temel kanıtının manifestte stale SHA taşımasını reddetmiş ve üretici sırası FAIL'i 2abcf853 ile korunmuştur. Final üretici sırası governed preflight writer ardından manifest/SHA256SUMS son üretimdir; yeni exact committe tüm kapılar tekrarlanır.",
    "V5 30 sayfa renderında sayfa 12/14 ilk çoklu önizleme yorumu görüntüleyici kırpmasını dosya kusuru sanmış ve d1e0b803 ile tarihsel yanlış pozitif olarak korunmuştur. Ayrı özgün çözünürlük/piksel doğrulaması dört sütun başlıklarını eksiksiz bulmuş ve 30/30 görsel PASS vermiştir. Ana karar sicilinin bağlayıcı saydığı halde kaynakta bulunmayan ADR-067 gerçek kayıt bütünlüğü FAIL'i f1590772 ile korunmuş; DEC-084 ve Migrasyon 38'in mevcut claim rezervasyonu gerçeğinden ADR geri kurulmuş, kesintisiz ADR numarası ve bağlayıcı referans eşliği fail-closed kapıya alınmıştır.",
    "V5 PDF Poppler renderı 26 sayfa vermiş; ilk turda sayfa 4/5 Yerel makine durum kodlarının token ortasında bölünmesi okunabilirlik FAIL olmuş ve 6d94ab9e ile korunmuştur. Kanonik makine değerleri değişmeden görünür DOCX/PDF durumları alt çizgi yerine anlamsal boşluklarla sarılmış; final DOCX 30/30 ve PDF 26/26, toplam 56/56 özgün çözünürlük sayfa taşma, örtüşme, kırpılma, font/glyph, tablo, footer ve sayfa numarası kusuru olmadan PASS vermiştir.",
    "Final-freeze render aracının desteklenmeyen argümanlarla ilk çağrısı belge testi başlamadan durmuş ve 764e856b boş reddedilmiş checkpointiyle korunmuştur. Doğru PATH tabanlı retry DOCX 30 ve PDF 26 sayfayı eksiksiz üretmiştir; bu çağrı reddi ürün veya belge kusuru değildir.",
    "Final-freeze2 çoklu önizlemesindeki DOCX header/footer ve PDF çift sayfa footer kırpması tam sayfa readback ile yanlış pozitiftir. PDF karar dizinindeki exact yolların karakter ortasından sarılması gerçek FAIL olup 17ad92d0 ile korunmuştur; karar/ADR yol hücreleri exact metni değiştirmeyen ayraç-sonrası sıfır-genişlikli kırma noktalarıyla yeniden üretilir.",
    "Final-freeze3 makine kapısı 30 DOCX ve 27 PDF sayfayı eksiksiz bulmuş, ancak U+200B ReportLab token bölmesini engellememiştir. PDF karar/ADR yollarındaki ayraç dışı sarım a0d9df42 ile korunmuş; yol üreticisi yalnız /, -, _ sonrasında en çok 48 karakterlik deterministik satırlar üretir. Final-freeze4 DOCX 30/30 ve PDF 27/27, toplam 57/57 tam tek-sayfa özgün çözünürlük QA PASS vermiştir.",
    "0099e39e yanlış ana kaynak kökü mutation assessment çağrısını fail-closed reddeder; exact Bronze retry 109 değişen yol/21 hedef test ve analysis 109 yol PASS vermiştir. P2 kayıt bütünlüğü sertleştirmesi assessment sourceCommit değerini canlı release provenance HEAD, baselineCommit değerini doğrulanmış harici baseline pointer HEAD ile exact bağlar. Eksik veya drift kimlik yedi tüketicide reddedilir; odaklı 2 dosya/8 test PASS, yeni exact hedefli/tam/bütünlük ve kurulu UAT zinciri pendingdir.",
    "27.08.2026 Bronze 52 kapanışında Bronze 51 predecessor PASS'tir. d68fd2a4 invocation-only; 3976994d, fb8683dc, a6020cb4 ve 61f09ed5 gerçek retlerdir. ASCII char düzeltmesi 2 dosya/18 test, bütün TICARI-052 marker retryı 1.254 kontrol/87 dosya/61 iş/241 kural PASS'tir. Bronze 27.08.2026.52 tek kez tahsis edilmiştir; yeni exact test, kaynak bütünlüğü, belge QA, paket ve N-to-N+1 kurulu UAT kapanmadan teslim yoktur.",
    "f4f84896 exact Bronze 52 kaynak turunda ana kaynak ve Bronze/Silver/Gold temiz eşitlik, salt-okunur preflight, 88 değişen yol/19 hedef test assessment/analysis ve hedefli 19 dosya/191 test PASS vermiştir. Filtresiz tur 399 dosya/2.483 test PASS iken yalnız PPK-015 current source inventory SHA ratchet'i eski kalmış ve gerçek FAIL 8ea2dfe1 ile korunmuştur. Canlı sınır 18 bölge/590 dosya/0 bulgu/2 adapter/3 amaç/4 yalnız-yerel taşıma dosyasıdır; kaynak özeti f54e3f302649af67ed6d028e66673eea68b0d58c2ba43c912c1ccb7534babe98 değerine ilerlemiş, ağ yetkisi değişmemiştir. Ratchet ve bağımlı kayıtlar eşlenir; yeni exact zincir zorunludur.",
    "İlk 32-K PPK-015 contract retryı, latest migration 121 bilgisi doğru kalırken zorunlu exact küçük harfli migration 117 tarihsel ayırıcısının kaybını 43 kontrolde 1 gerçek FAIL olarak bulmuş ve 24e6bd71 ile korumuştur. Ayırıcı geri kurulur; canlı sınır ve ağ yetkisi değişmez, taze contract/runtime/ticari ve bütün exact zincir zorunludur.",
    "PPK-015 tarihsel ayırıcı düzeltmesi sonrası retry contract 43/43, runtime 10/10, iki odaklı test dosyası 23/23 ve ticari temel 1.254 kontrol PASS vermiştir. 24e6bd71 gerçek ret olarak korunur; bu odaklı PASS yeni exact committe kanonik hedefli/tam regresyon, kaynak bütünlüğü, paket ve kurulu UAT yerine geçmez.",
    "Final-freeze6 P2 belge QA sonucunda önceki onayla byte-exact aynı 25 sayfa korunmuş, değişen DOCX 1 ve 9-19 ile PDF 1 ve 9-27 sayfaları üç bağımsız denetimde 32/32 PASS bulunmuştur. Toplam DOCX 30/30 ve PDF 27/27, yani 57/57 sayfa; taşma, örtüşme, kırpılma, font/glyph, tablo, footer, marj, sayfa numarası ve güvenli ayraç dışı token bölünmesi olmadan PASS'tir.",
    "Bronze 27.08.2026.52 güncel master belge QA turunda DOCX 30/30 ve PDF 28/28, toplam 58/58 sayfa özgün çözünürlükte PASS vermiştir. Taşma, örtüşme, kırpılma, token-ortası sarım, font/glyph, tablo, header/footer, marj veya sayfa numarası kusuru yoktur; exact ürün, paket ve kurulu tam UI UAT kapıları ayrıca zorunludur.",
    "PPK-015 ret/retry kayıtlarıyla yeniden üretilen güncel master DOCX 31/31 ve PDF 28/28, toplam 59/59 sayfada görsel QA PASS vermiştir. Bütün sayfalar temas görünümünde; değişiklik yoğun 7-9 ve tablo yoğun 25-31 ayrıca özgün çözünürlükte taşma, örtüşme, kırpılma, token-ortası sarım, font/glyph, tablo, header/footer, marj veya sayfa numarası kusuru olmadan PASS'tir. Exact ürün, paket ve kurulu UAT kapıları ayrıca zorunludur.",
]


def package_evidence(item: dict) -> str:
    missing = item.get("missingEvidence") or []
    return ", ".join(missing[:8]) + (f" (+{len(missing) - 8})" if len(missing) > 8 else "") if missing else "Yok/yerel tamamlandı"


def package_status_reason(item: dict) -> str:
    if item.get("status") == "COMPLETED":
        return item.get("closureReason") or "Kapalı; yerel kabul zinciri ve kalıcı receipt tamamlandı."
    return item.get("openReason") or "Açık kalma nedeni kaydedilmemiş — fail-closed belge drifti."


def build_markdown() -> str:
    lines = [
        "# ParsYuva Aile Yaşam Merkezi — Güncel Karar, Kural ve İş Akışı Sicili",
        "",
        f"- Belge sürümü: **{VERSION}**",
        f"- Tarih: **{AS_OF}**",
        f"- Görünür ürün sürümü: **{ledger['release']}**",
        f"- Kaynak HEAD: `{git_head()}`",
        "- Statü: **ACTIVE_CURRENT_MASTER_REFERENCE**",
        "- Kararlar: **DEC-250–DEC-276**",
        "",
        "> Bu sürüm geçmiş PDF/DOCX ve build kapanış belgelerinin üzerine yazmaz. Yerel PASS ile dış kabul kanıtını ayırır; NOT_RUN/PARTIAL/BLOCKED sonuçlarını tamamlanmış göstermez.",
        "",
        "## 1. Denetim sonucu",
        "",
        f"- Aktif repo Word/PDF tarihsel taraması: **{len(word_pdf_scan)} dosya / {historical_pairs} çift / {readable_count} okunabilir**.",
        f"- `C:\\PPT\\AYM` tüm belge türü taraması: **{full_document_audit['documentFileCount']} dosya / {full_document_audit['readableCount']} okunabilir / {full_document_audit['unreadableCount']} sorun**.",
        f"- Office/RTF/PDF: **{full_document_audit['officeAndPdfCount']}**; benzersiz içerik hash'i: **{full_document_audit['uniqueContentHashCount']}**; tekrar kopya: **{full_document_audit['duplicateFileCount']}**.",
        "- Build209–228 master çiftleri ve eski Bronze aktif referans çifti tarihsel olarak korunmuştur.",
        f"- Karar dosyası: **{len(decision_files)}**; ADR: **{len(adr_files)}**; security/threat belgesi: **{len(threat_files)}**.",
        f"- Mevcut tam belge/config/kanıt envanteri: **{document_index['documentCount']}** (yeni sürümden önceki indeks).",
        "",
        "## 2. Yetki ve öncelik",
        "",
        "1. Kanonik kural sicili ve Proje Anayasası",
        "2. Aktif yönetişim, kullanıcı karar ve kabul edilmiş kapsam sicilleri",
        "3. Güncel birleşik sicil ve aktif çalışma belgeleri",
        "4. DEC/ADR/threat model ve makine okunur scope/inventory",
        "5. Kaynak kod, test ve üretilmiş kanıt",
        "6. Tarihsel build belgeleri (yalnız kendi zamanlarının kanıtı)",
        "",
        "## 3. Kapsam ve kural özeti",
        "",
        f"- Gereksinim: **{requirements['requirementCount']}** — COMPLETE {requirement_counts.get('COMPLETE', 0)}, PARTIAL {requirement_counts.get('PARTIAL', 0)}, FOUNDATION_STARTED {requirement_counts.get('FOUNDATION_STARTED', 0)}, NOT_IMPLEMENTED {requirement_counts.get('NOT_IMPLEMENTED', 0)}.",
        f"- Kural sicili: **{rules['id']}**, toplam {rules['ruleCount']}, aktif {rules['activeRuleCount']}, superseded {rules['supersededRuleCount']}, SHA-256 `{rules['rulesSha256']}`.",
        f"- Kullanıcı karar defteri: **{user_decisions['decisionCount']}** toplam kullanıcı kararı; **{decision_counts.get('ACTIVE', 0)}** aktif, **{decision_counts.get('SUPERSEDED', 0)}** superseded.",
        "",
        "## 4. İş akışları",
        "",
    ]
    for title, detail in WORKFLOWS:
        lines += [f"### {title}", "", detail, ""]
    lines += ["## 5. Paket iş listesi — açık/kapalı/neden", "", "| Paket | Resmî durum | Yerel durum | Requirement PASS | Açık kalma nedeni | Eksik kanıt |", "|---|---|---|---:|---|---|"]
    for item in roadmap["packages"]:
        lines.append(
            f"| {item['step']} | {item['status']} | {item.get('localImplementationStatus') or '-'} | {'EVET' if item.get('countsAsRequirementPass') else 'HAYIR'} | {package_status_reason(item)} | {package_evidence(item)} |"
        )
    lines += [
        "",
        "## 6. 34-L yerel kapanış gerçeği",
        "",
        f"- Boundary: {closure['validation']['localPackageBoundaries']['status']} / {closure['validation']['localPackageBoundaries']['checks']} kontrol.",
        f"- Contract: {closure['validation']['localPackageContracts']['status']} / {closure['validation']['localPackageContracts']['checks']} kontrol.",
        f"- Runtime: {closure['validation']['localPackageRuntimes']['status']} / {closure['validation']['localPackageRuntimes']['checks']} kontrol.",
        f"- Full regression: {closure['validation']['fullRegression']['status']} / {closure['validation']['fullRegression']['files']} dosya / {closure['validation']['fullRegression']['tests']} test.",
        f"- Production build: {closure['validation']['productionBuilds']['status']} / {closure['validation']['productionBuilds']['workspaces']} workspace.",
        f"- Güncel Core Service companion: {windows_scope['validation']['currentCoreServiceCompanionEvidence']['status']} / {windows_scope['validation']['currentCoreServiceCompanionEvidence']['fullRegression']['files']} dosya / {windows_scope['validation']['currentCoreServiceCompanionEvidence']['fullRegression']['tests']} test / {windows_scope['validation']['currentCoreServiceCompanionEvidence']['packagedLaunches']} normal paket açılışı.",
        f"- Güncel dağıtım imzası: {windows_scope['validation']['currentCoreServiceCompanionEvidence']['executableSignature']}; yerel imzasız ParsYuva-Bronze-20.08.2026.37.exe üretildi ve aynı win-unpacked paketinin iki ardışık açılışı PASS verdi. Production sertifikası ve yükseltilmiş kurulu yaşam döngüsü PASS olmadığı için ticari dağıtım hazır sayılmaz.",
        "- Buna rağmen allRoadmapPackagesAccepted=false, requirementsClosed=false ve countsAsRequirementPass=false.",
        "",
        "## 7. Dış bağımlılıklar ve neden açık",
        "",
    ]
    for title, detail in EXTERNAL_DEPENDENCIES:
        lines += [f"- **{title}:** {detail}"]
    lines += ["", "## 8. Belge sapmaları ve yapılan düzeltmeler", ""]
    lines += [f"- {item}" for item in DRIFT_FIXES]
    lines += ["", "## 9. Görsel kimlik ve erişilebilirlik", ""]
    lines += [
        f"- Logo: `{visual['brandMark']['path']}`, `{visual['brandMark']['style']}`, SHA-256 `{visual['brandMark']['sha256']}`, {visual['brandMark']['width']}×{visual['brandMark']['height']}, şeffaf arka plan.",
        f"- Tipografi: large title {visual['typography']['largeTitlePx']} px; title1 {visual['typography']['title1Px']} px; title2 {visual['typography']['title2Px']} px; body {visual['typography']['bodyPx']} px; control {visual['typography']['controlPx']} px; minimum {visual['typography']['minimumPx']} px.",
    ]
    for channel, palette in visual["releaseChannelNavigationColors"].items():
        lines.append(f"- {channel}: text {palette['text']}, strong {palette['strong']}, icon {palette['icon']}, edge {palette['edge']}.")
    lines += ["", "## 10. Tüm belge türü denetimi", ""]
    for extension, count in full_document_audit["extensionCounts"].items():
        lines.append(f"- `{extension}`: {count}")
    lines += [f"- Bulunmayan Office türleri: {', '.join(full_document_audit['absentOfficeFormats']) or 'Yok'}.", "- Tam yol/hash/okuma sonucu: `artifacts/manifests/ALL_PROJECT_DOCUMENT_FORMAT_AUDIT.json`."]
    lines += ["", "## 11. DEC karar dizini", ""]
    for path in decision_files:
        lines.append(f"- `{path.stem.split('-')[0]}-{path.stem.split('-')[1]}` — {first_heading(path)} — {status_from_markdown(path)} — `{path.relative_to(ROOT).as_posix()}`")
    lines += ["", "## 12. ADR dizini", ""]
    for path in adr_files:
        lines.append(f"- `{path.stem.split('-')[0]}-{path.stem.split('-')[1]}` — {first_heading(path)} — `{path.relative_to(ROOT).as_posix()}`")
    lines += ["", "## 13. Kanonik kurallar — eksiksiz", ""]
    for rule in rules["rules"]:
        lines.append(f"- **{rule['id']} [{rule['state']}]** — {rule['text']}")
    lines += ["", "## 14. Aktif repo Word/PDF tarihsel denetim envanteri", ""]
    for row in word_pdf_scan:
        extra = f", {row.get('pages', '?')} sayfa" if row.get("pages") is not None else ""
        lines.append(f"- `{row['name']}` — {row['type'].upper()} — {row['bytes']} bayt{extra} — {'OKUNABİLİR' if row['readable'] else 'HATA'}")
    lines += ["", "## 15. Kapanış sınırı", "", "Bu belge canlı kaynak gerçeğini toplar; build kapanışı, kanal terfisi, sertifika, hukuk görüşü veya gerçek cihaz/sağlayıcı UAT belgesi değildir. Tarihsel kanıtlar değişmeden kalır.", ""]
    return "\n".join(lines)


SOURCE.parent.mkdir(parents=True, exist_ok=True)
SOURCE.write_bytes(build_markdown().encode("utf-8"))


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=8.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_row_pagination(row, *, repeat_header=False):
    row_properties = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    row_properties.append(cannot_split)
    if repeat_header:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        row_properties.append(header)


def set_table_fixed(table, widths: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_doc_heading(document, text: str, level=1):
    p = document.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_doc_paragraph(document, text: str, *, bold_prefix=None, italic=False):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        second = p.add_run(text[len(bold_prefix):])
        set_run_font(second)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    return p


def add_doc_table(document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=7.8):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_row_pagination(table.rows[0], repeat_header=True)
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, BRONZE)
        set_cell_text(cell, value, bold=True, color="FFFFFF", size=font_size)
    for values in rows:
        row = table.add_row()
        set_row_pagination(row)
        cells = row.cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value, size=font_size)
    set_table_fixed(table, widths)
    return table


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.95)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, BRONZE),
        ("Heading 2", 13, 14, 7, BRONZE),
        ("Heading 3", 11, 10, 5, GREEN),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("PARSYUVA AİLE YAŞAM MERKEZİ  •  GÜNCEL MASTER DOKÜMANTASYON"), size=8, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run(f"{VERSION}  •  Sayfa "), size=8, color=MUTED)
    add_page_field(footer)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(1.25))
        logo_shape._inline.docPr.set("title", "ParsYuva logosu")
        logo_shape._inline.docPr.set("descr", "Sıcak bronz tonlarda Anadolu parsı başı biçimindeki ParsYuva logosu")
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run("GÜNCEL ÇALIŞMA REFERANSI"), size=10, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("PARSYUVA AİLE YAŞAM MERKEZİ"), size=24, color=BRONZE, bold=True)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(16)
    set_run_font(sub2.add_run("Karar • Kural • Mimari • İş Akışı • Açık İş • Kanıt Sicili"), size=12.5, color=GREEN)
    add_doc_table(doc, ["Alan", "Güncel değer"], [
        ["Belge sürümü", VERSION], ["Tarih", AS_OF], ["Ürün sürümü", ledger["release"]],
        ["Kaynak HEAD", git_head()], ["Kural sicili", f"{rules['id']} / {rules['activeRuleCount']} aktif / {rules['rulesSha256']}"],
        ["Durum", "Aktif çalışma referansı; build kapanışı veya kanal terfisi değildir"],
    ], [1900, 7460], 8.4)
    add_doc_paragraph(doc, "Tarihsel Build209–228 ve eski Bronze master belgeleri değiştirilmeden korunmuştur. Yerel PASS, dış kabul kanıtı değildir; NOT_RUN/PARTIAL/BLOCKED tamamlandı gösterilmez.", italic=True)

    doc.add_page_break()
    add_doc_heading(doc, "1. Yönetici özeti", 1)
    add_doc_paragraph(doc, f"Denetimde {len(word_pdf_scan)} mevcut Word/PDF dosyasının tamamı ayrıştırıldı; {readable_count} dosya okunabilir bulundu. Tarihsel çiftler korunurken güncel master sürüm ayrı adla oluşturuldu.")
    summary_rows = [
        ["Kabul edilmiş gereksinim", str(requirements["requirementCount"])],
        ["COMPLETE / PARTIAL / FOUNDATION / NOT_IMPLEMENTED", f"{requirement_counts.get('COMPLETE',0)} / {requirement_counts.get('PARTIAL',0)} / {requirement_counts.get('FOUNDATION_STARTED',0)} / {requirement_counts.get('NOT_IMPLEMENTED',0)}"],
        ["Kanonik kural", f"{rules['ruleCount']} toplam / {rules['activeRuleCount']} aktif / {rules['supersededRuleCount']} superseded"],
        ["Karar / ADR / threat dosyası", f"{len(decision_files)} / {len(adr_files)} / {len(threat_files)}"],
        ["İş paketi", f"{roadmap['packageCount']} toplam; 3 resmî COMPLETED, 23 açık/acceptance blocked"],
        ["Silver", "BLOCKED"],
    ]
    add_doc_table(doc, ["Gösterge", "Canlı sonuç"], summary_rows, [3000, 6360], 8.5)

    add_doc_heading(doc, "2. Belge yetkisi ve tarihsel koruma", 1)
    for text in [
        "Kanonik kural sicili ve Proje Anayasası en üst makine-okunur otoritedir.",
        "Aktif governance, kullanıcı karar ve kabul edilmiş kapsam sicilleri çalışma durumunu belirler.",
        "Bu güncel master, aktif Markdown/config/DEC/ADR/threat/source/test kanıtlarını birleştirir.",
        "Build209–228 ve eski Bronze Word/PDF dosyaları kendi zamanlarının tarihsel delilidir; aktif davranışı geçersiz kılamaz ve üzerine yazılmaz.",
    ]:
        add_doc_paragraph(doc, text)

    add_doc_heading(doc, "3. Kanonik kapsam ve kural tabanı", 1)
    add_doc_paragraph(doc, f"Kapsam sicilinde {requirements['requirementCount']} gereksinim vardır. Dağılım: COMPLETE {requirement_counts.get('COMPLETE',0)}, PARTIAL {requirement_counts.get('PARTIAL',0)}, FOUNDATION_STARTED {requirement_counts.get('FOUNDATION_STARTED',0)}, NOT_IMPLEMENTED {requirement_counts.get('NOT_IMPLEMENTED',0)}.")
    add_doc_paragraph(doc, f"Kural sicili {rules['id']}: {rules['ruleCount']} toplam, {rules['activeRuleCount']} aktif, {rules['supersededRuleCount']} superseded. SHA-256 {rules['rulesSha256']}.")
    add_doc_paragraph(doc, "Tamamlanma zinciri karar → domain → şema/migration → use-case → repository → PEP/UoW → API/IPC → UI/menü → hedefli test → belge → kanıttır. Zincirin herhangi bir halkası eksikse requirement tamamlanmaz.")

    add_doc_heading(doc, "4. Güncel altyapı ve mimari", 1)
    add_doc_table(doc, ["Katman", "Güncel altyapı gerçeği"], [[a, b] for a, b in INFRASTRUCTURE], [2100, 7260], 8.1)

    add_doc_heading(doc, "5. Bağlayıcı iş akışları", 1)
    add_doc_table(doc, ["İş akışı", "Bağlayıcı yürütme"], [[title, detail] for title, detail in WORKFLOWS], [2300, 7060], 8.4)

    add_doc_heading(doc, "6. Paket iş listesi — açık/kapalı/neden", 1)
    add_doc_paragraph(doc, "Kullanıcı kuralı gereği her açık paket; yerel olarak neyin bulunduğunu, neden açık kaldığını, eksik kanıtı ve requirement PASS sayılıp sayılmadığını birlikte gösterir.")
    package_rows = []
    for item in roadmap["packages"]:
        package_rows.append([
            item["step"], visible_status(item["status"]), visible_status(item.get("localImplementationStatus")),
            "EVET" if item.get("countsAsRequirementPass") else "HAYIR",
            package_status_reason(item),
        ])
    add_doc_table(doc, ["Paket", "Resmî", "Yerel", "PASS", "Açık kalma nedeni"], package_rows, [650, 1050, 1700, 650, 5310], 6.9)

    add_doc_heading(doc, "7. Yerel doğrulama ve kabul sınırı", 1)
    validation = closure["validation"]
    validation_rows = [
        ["Boundary", visible_status(validation["localPackageBoundaries"]["status"]), str(validation["localPackageBoundaries"]["checks"])],
        ["Contract", visible_status(validation["localPackageContracts"]["status"]), str(validation["localPackageContracts"]["checks"])],
        ["Runtime", visible_status(validation["localPackageRuntimes"]["status"]), str(validation["localPackageRuntimes"]["checks"])],
        ["Targeted", visible_status(validation["targeted"]["status"]), f"{validation['targeted']['files']} dosya / {validation['targeted']['tests']} test"],
        ["Full regression", visible_status(validation["fullRegression"]["status"]), f"{validation['fullRegression']['files']} dosya / {validation['fullRegression']['tests']} test"],
        ["Root typecheck", validation["rootTypecheck"], "Kaynak tipi"],
        ["Production builds", visible_status(validation["productionBuilds"]["status"]), f"{validation['productionBuilds']['workspaces']} workspace"],
        ["Requirement kabul", "HAYIR", "requirementsClosed=false / countsAsRequirementPass=false"],
    ]
    add_doc_table(doc, ["Kapı", "Sonuç", "Kanıt"], validation_rows, [2100, 1400, 5860], 8.1)

    add_doc_heading(doc, "8. Dış bağımlılıklar ve neden açık", 1)
    add_doc_table(doc, ["Dış bağımlılık", "Açık kalma nedeni"], [[title, detail] for title, detail in EXTERNAL_DEPENDENCIES], [2300, 7060], 8.4)

    add_doc_heading(doc, "9. Installer ve çalıştırma gerçeği", 1)
    add_doc_paragraph(doc, "Güncel kanal program hedefi legacy kökün dışındaki C:\\Program Files\\PPT\\ParsYuva-<Kanal> kardeş dizinidir; ana dosya ParsYuva-<Kanal>.exe, kısayol ParsYuva <Kanal>, AppData kökü ParsYuva/<Kanal> ve teslim EXE adı ParsYuva-<Kanal>-GG.AA.YYYY.NN.exe biçimindedir. Otomatik legacy veri migration veya silme yoktur. ParsYuva-Bronze-20.08.2026.37.exe tarihsel yerel test installerı ile aynı win-unpacked paketin iki açılış kaydı güncel N->N+1 kabul kanıtı değildir. Yükseltilmiş gerçek kurulum yaşam döngüsü PASS olmadıkça ve Production Authenticode sertifikası ile temiz harici Windows makinesi kanıtı tamamlanmadıkça 'ticari dağıtıma hazır' iddiası kurulmaz.")

    add_doc_heading(doc, "10. Görsel kimlik ve erişilebilirlik", 1)
    palette_rows = [[channel, value["text"], value["strong"], value["icon"], value["edge"]] for channel, value in visual["releaseChannelNavigationColors"].items()]
    add_doc_table(doc, ["Kanal", "Text", "Strong", "Icon", "Edge"], palette_rows, [1200, 2040, 2040, 2040, 2040], 8.0)
    add_doc_paragraph(doc, f"Logo {visual['brandMark']['width']}×{visual['brandMark']['height']} şeffaf PNG, stil {visual['brandMark']['style']}, SHA-256 {visual['brandMark']['sha256']}. Body {visual['typography']['bodyPx']} px; control {visual['typography']['controlPx']} px; minimum {visual['typography']['minimumPx']} px. Gerçek ekran/Narrator/büyütme/UAT ayrıca gereklidir.")

    add_doc_heading(doc, "11. Belge sapmaları ve düzeltmeler", 1)
    for item in DRIFT_FIXES:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.keep_together = True
        set_run_font(p.add_run(item), size=9.5)

    add_doc_heading(doc, "12. Tüm belge türü denetimi", 1)
    add_doc_table(doc, ["Uzantı", "Dosya"], [[extension, str(count)] for extension, count in full_document_audit["extensionCounts"].items()], [2800, 6560], 8.2)
    add_doc_paragraph(doc, f"Kök taramada {full_document_audit['documentFileCount']} belge/config/metin dosyası bulundu; {full_document_audit['readableCount']} okunabilir, {full_document_audit['unreadableCount']} sorunlu. Tam yol ve SHA-256 listesi artifacts/manifests/ALL_PROJECT_DOCUMENT_FORMAT_AUDIT.json içindedir.")

    add_doc_heading(doc, "13. DEC karar dizini — eksiksiz", 1)
    decision_rows = [[f"DEC-{decision_number(path):03d}", first_heading(path), status_from_markdown(path), path.relative_to(ROOT).as_posix()] for path in decision_files]
    add_doc_table(doc, ["ID", "Karar", "Durum", "Dosya"], decision_rows, [850, 4140, 1500, 2870], 6.5)

    add_doc_heading(doc, "14. ADR dizini — eksiksiz", 1)
    adr_rows = [[f"ADR-{adr_number(path):03d}", first_heading(path), path.relative_to(ROOT).as_posix()] for path in adr_files]
    add_doc_table(doc, ["ID", "Mimari karar", "Dosya"], adr_rows, [900, 5000, 3460], 6.8)

    add_doc_heading(doc, f"15. Kanonik kurallar — {rules['ruleCount']} kayıt", 1)
    rule_rows = [[rule["id"], rule["state"], rule["text"]] for rule in rules["rules"]]
    add_doc_table(doc, ["ID", "Durum", "Kural"], rule_rows, [850, 1100, 7410], 7.0)

    add_doc_heading(doc, "16. Aktif repo Word/PDF tarihsel envanter denetimi", 1)
    audit_rows = [[row["name"], row["type"].upper(), str(row["bytes"]), str(row.get("pages", "-")), "OK" if row["readable"] else "HATA"] for row in word_pdf_scan]
    add_doc_table(doc, ["Dosya", "Tür", "Bayt", "Sayfa", "Okuma"], audit_rows, [5000, 800, 1300, 850, 1410], 7.0)
    add_doc_paragraph(doc, "Yeni sürüm, tarihsel dosyaların yerine geçmez; yalnız aktif çalışma gerçeğini tek yerde toplar.")

    doc.core_properties.title = "ParsYuva Aile Yaşam Merkezi Güncel Master Dokümantasyon"
    doc.core_properties.subject = "Karar, kural, mimari, iş akışı, açık iş ve kanıt sicili"
    doc.core_properties.author = "ParsYuva"
    doc.core_properties.keywords = "ParsYuva Aile Yaşam Merkezi, governance, güvenlik, kurumsallaşma, iş akışı, karar, kural"
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def pdf_escape(text: str) -> str:
    return html.escape(str(text)).replace("\n", "<br/>")


def pdf_path_breaks(value: str, max_chars: int = 48) -> str:
    """Keep exact path text and add deterministic line breaks only after separators."""
    lines: list[str] = []
    current = ""
    for part in re.split(r"(?<=[/_-])", str(value)):
        if current and len(current) + len(part) > max_chars:
            lines.append(current)
            current = part
        else:
            current += part
    if current:
        lines.append(current)
    return "\n".join(lines)


def build_pdf():
    arial = Path("C:/Windows/Fonts/arial.ttf")
    arial_bold = Path("C:/Windows/Fonts/arialbd.ttf")
    font = "Helvetica"
    bold = "Helvetica-Bold"
    if arial.exists() and arial_bold.exists():
        pdfmetrics.registerFont(TTFont("ArialCurrent", str(arial)))
        pdfmetrics.registerFont(TTFont("ArialCurrent-Bold", str(arial_bold)))
        font, bold = "ArialCurrent", "ArialCurrent-Bold"
    styles = getSampleStyleSheet()
    body = ParagraphStyle("BodyCurrent", parent=styles["BodyText"], fontName=font, fontSize=8.6, leading=11, spaceAfter=5, textColor=colors.HexColor(f"#{INK}"))
    small = ParagraphStyle("SmallCurrent", parent=body, fontSize=6.7, leading=8.2, spaceAfter=2)
    h1 = ParagraphStyle("H1Current", parent=styles["Heading1"], fontName=bold, fontSize=14, leading=17, textColor=colors.HexColor(f"#{BRONZE}"), spaceBefore=10, spaceAfter=7)
    h2 = ParagraphStyle("H2Current", parent=styles["Heading2"], fontName=bold, fontSize=10.5, leading=13, textColor=colors.HexColor(f"#{GREEN}"), spaceBefore=7, spaceAfter=4)
    title_style = ParagraphStyle("TitleCurrent", parent=styles["Title"], fontName=bold, fontSize=23, leading=27, alignment=TA_CENTER, textColor=colors.HexColor(f"#{BRONZE}"), spaceAfter=5)
    subtitle_style = ParagraphStyle("SubtitleCurrent", parent=body, fontName=font, fontSize=12, leading=15, alignment=TA_CENTER, textColor=colors.HexColor(f"#{GREEN}"), spaceAfter=14)

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont(font, 7)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawString(0.65 * inch, 0.36 * inch, f"ParsYuva Aile Yaşam Merkezi • {VERSION}")
        canvas.drawRightString(7.85 * inch, 0.36 * inch, f"Sayfa {doc.page}")
        canvas.restoreState()

    def p(text: str, style=body):
        return Paragraph(pdf_escape(text), style)

    def pdf_table(headers, rows, widths, font_size=7.0):
        data = [[Paragraph(pdf_escape(value), ParagraphStyle("TH", parent=small, fontName=bold, fontSize=font_size, leading=font_size + 1.5, textColor=colors.white)) for value in headers]]
        for row in rows:
            data.append([Paragraph(pdf_escape(value), ParagraphStyle("TD", parent=small, fontSize=font_size, leading=font_size + 1.5)) for value in row])
        table = Table(data, colWidths=widths, repeatRows=1, splitByRow=1, splitInRow=0, longTableOptimize=False)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BRONZE}")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8B8A8")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{BRONZE_LIGHT}")]),
        ]))
        return table

    story = [Spacer(1, 0.22 * inch)]
    if LOGO.exists():
        logo = Image(str(LOGO), width=1.05 * inch, height=1.05 * inch)
        logo.hAlign = "CENTER"
        story += [logo, Spacer(1, 6)]
    story += [
        Paragraph("GÜNCEL ÇALIŞMA REFERANSI", ParagraphStyle("Kicker", parent=body, alignment=TA_CENTER, fontName=bold, fontSize=9, textColor=colors.HexColor(f"#{GOLD}"))),
        Paragraph("PARSYUVA AİLE YAŞAM MERKEZİ", title_style),
        Paragraph("Karar • Kural • Mimari • İş Akışı • Açık İş • Kanıt Sicili", subtitle_style),
        pdf_table(["Alan", "Güncel değer"], [
            ["Belge sürümü", VERSION], ["Tarih", AS_OF], ["Ürün sürümü", ledger["release"]], ["Kaynak HEAD", git_head()],
            ["Kural sicili", f"{rules['id']} / {rules['activeRuleCount']} aktif / {rules['rulesSha256']}"],
            ["Durum", "Aktif çalışma referansı; build kapanışı veya kanal terfisi değildir"],
        ], [1.45 * inch, 5.55 * inch], 8),
        Spacer(1, 8), p("Tarihsel Build209–228 ve eski Bronze master belgeleri değiştirilmeden korunmuştur. Yerel PASS, dış kabul kanıtı değildir; NOT_RUN/PARTIAL/BLOCKED tamamlandı gösterilmez."),
        PageBreak(), Paragraph("1. Yönetici özeti", h1),
        p(f"Denetimde {len(word_pdf_scan)} mevcut Word/PDF dosyasının tamamı ayrıştırıldı; {readable_count} dosya okunabilir bulundu. Tarihsel çiftler korunurken güncel master sürüm ayrı adla oluşturuldu."),
        pdf_table(["Gösterge", "Canlı sonuç"], [
            ["Kabul edilmiş gereksinim", str(requirements["requirementCount"])],
            ["COMPLETE / PARTIAL / FOUNDATION / NOT_IMPLEMENTED", f"{requirement_counts.get('COMPLETE',0)} / {requirement_counts.get('PARTIAL',0)} / {requirement_counts.get('FOUNDATION_STARTED',0)} / {requirement_counts.get('NOT_IMPLEMENTED',0)}"],
            ["Kanonik kural", f"{rules['ruleCount']} toplam / {rules['activeRuleCount']} aktif / {rules['supersededRuleCount']} superseded"],
            ["Karar / ADR / threat", f"{len(decision_files)} / {len(adr_files)} / {len(threat_files)}"],
            ["İş paketi", f"{roadmap['packageCount']} toplam; 3 resmî COMPLETED, 23 açık/acceptance blocked"], ["Silver", "BLOCKED"],
        ], [2.4 * inch, 4.6 * inch], 8),
        Paragraph("2. Belge yetkisi ve tarihsel koruma", h1),
    ]
    for text in [
        "Kanonik kural sicili ve Proje Anayasası en üst makine-okunur otoritedir.",
        "Aktif governance, kullanıcı karar ve kabul edilmiş kapsam sicilleri çalışma durumunu belirler.",
        "Bu güncel master, aktif Markdown/config/DEC/ADR/threat/source/test kanıtlarını birleştirir.",
        "Build209–228 ve eski Bronze Word/PDF dosyaları kendi zamanlarının tarihsel delilidir; aktif davranışı geçersiz kılamaz ve üzerine yazılmaz.",
    ]:
        story.append(p(text))
    story += [Paragraph("3. Kanonik kapsam ve kural tabanı", h1), p(f"Kapsam sicilinde {requirements['requirementCount']} gereksinim vardır. Dağılım: COMPLETE {requirement_counts.get('COMPLETE',0)}, PARTIAL {requirement_counts.get('PARTIAL',0)}, FOUNDATION_STARTED {requirement_counts.get('FOUNDATION_STARTED',0)}, NOT_IMPLEMENTED {requirement_counts.get('NOT_IMPLEMENTED',0)}."), p(f"Kural sicili {rules['id']}: {rules['ruleCount']} toplam, {rules['activeRuleCount']} aktif, {rules['supersededRuleCount']} superseded. SHA-256 {rules['rulesSha256']}.")]
    # Bu tablo tek bir kısa devam satırı için yeni sayfa üretmemeli. Biraz daha
    # sıkı tipografi, bütün mimari özetini aynı sayfada okunaklı biçimde tutar.
    story += [Paragraph("4. Güncel altyapı ve mimari", h1), pdf_table(["Katman", "Güncel altyapı gerçeği"], [[a, b] for a, b in INFRASTRUCTURE], [1.55 * inch, 5.45 * inch], 7.2), PageBreak(), Paragraph("5. Bağlayıcı iş akışları", h1)]
    for title_text, detail in WORKFLOWS:
        story.append(KeepTogether([Paragraph(pdf_escape(title_text), h2), p(detail)]))
    package_rows = [[item["step"], visible_status(item["status"]), visible_status(item.get("localImplementationStatus")), "EVET" if item.get("countsAsRequirementPass") else "HAYIR", package_status_reason(item)] for item in roadmap["packages"]]
    story += [PageBreak(), Paragraph("6. Paket iş listesi — açık/kapalı/neden", h1), p("Her açık paket yerel durum, açık kalma nedeni, eksik kanıt ve requirement PASS gerçeğiyle birlikte gösterilir."), pdf_table(["Paket", "Resmî", "Yerel", "PASS", "Açık kalma nedeni"], package_rows, [0.45*inch, 0.8*inch, 1.3*inch, 0.45*inch, 4.0*inch], 5.8)]
    validation = closure["validation"]
    validation_rows = [["Boundary", visible_status(validation["localPackageBoundaries"]["status"]), str(validation["localPackageBoundaries"]["checks"])], ["Contract", visible_status(validation["localPackageContracts"]["status"]), str(validation["localPackageContracts"]["checks"])], ["Runtime", visible_status(validation["localPackageRuntimes"]["status"]), str(validation["localPackageRuntimes"]["checks"])], ["Targeted", visible_status(validation["targeted"]["status"]), f"{validation['targeted']['files']} dosya / {validation['targeted']['tests']} test"], ["Full regression", visible_status(validation["fullRegression"]["status"]), f"{validation['fullRegression']['files']} dosya / {validation['fullRegression']['tests']} test"], ["Production builds", visible_status(validation["productionBuilds"]["status"]), f"{validation['productionBuilds']['workspaces']} workspace"], ["Requirement kabul", "HAYIR", "requirementsClosed=false / countsAsRequirementPass=false"]]
    story += [Paragraph("7. Yerel doğrulama ve kabul sınırı", h1), pdf_table(["Kapı", "Sonuç", "Kanıt"], validation_rows, [1.6*inch, 1.1*inch, 4.3*inch], 7.6), Paragraph("8. Dış bağımlılıklar ve neden açık", h1)]
    for title_text, detail in EXTERNAL_DEPENDENCIES:
        story.append(KeepTogether([Paragraph(pdf_escape(title_text), h2), p(detail)]))
    visual_table = pdf_table(["Kanal", "Text", "Strong", "Icon", "Edge"], [[c, v['text'], v['strong'], v['icon'], v['edge']] for c, v in visual['releaseChannelNavigationColors'].items()], [1.0*inch,1.5*inch,1.5*inch,1.5*inch,1.5*inch],7.5)
    story += [KeepTogether([Paragraph("9. Installer ve çalıştırma gerçeği", h1), p("Güncel kanal program hedefi legacy kökün dışındaki C:\\Program Files\\PPT\\ParsYuva-<Kanal> kardeş dizinidir; ana dosya ParsYuva-<Kanal>.exe, kısayol ParsYuva <Kanal>, AppData kökü ParsYuva/<Kanal> ve teslim EXE adı ParsYuva-<Kanal>-GG.AA.YYYY.NN.exe biçimindedir. Otomatik legacy veri migration veya silme yoktur. ParsYuva-Bronze-20.08.2026.37.exe tarihsel yerel test installerı güncel N->N+1 kabul kanıtı değildir. Yükseltilmiş gerçek kurulum yaşam döngüsü PASS olmadıkça ve Production Authenticode sertifikası ile temiz harici Windows makinesi kanıtı tamamlanmadıkça ticari dağıtım hazır sayılmaz.")]), KeepTogether([Paragraph("10. Görsel kimlik ve erişilebilirlik", h1), visual_table, p(f"Logo {visual['brandMark']['width']}×{visual['brandMark']['height']} şeffaf PNG; SHA-256 {visual['brandMark']['sha256']}. Body {visual['typography']['bodyPx']} px, control {visual['typography']['controlPx']} px, minimum {visual['typography']['minimumPx']} px.")])]
    story += [KeepTogether([Paragraph("11. Belge sapmaları ve düzeltmeler", h1), p(f"• {DRIFT_FIXES[0]}")])] + [p(f"• {item}") for item in DRIFT_FIXES[1:]]
    story += [Paragraph("12. Tüm belge türü denetimi", h1), pdf_table(["Uzantı", "Dosya"], [[extension, str(count)] for extension, count in full_document_audit["extensionCounts"].items()], [2.0*inch,5.0*inch], 8.0), p(f"Kök taramada {full_document_audit['documentFileCount']} belge/config/metin dosyası bulundu; {full_document_audit['readableCount']} okunabilir, {full_document_audit['unreadableCount']} sorunlu. Tam yol ve SHA-256 listesi artifacts/manifests/ALL_PROJECT_DOCUMENT_FORMAT_AUDIT.json içindedir.")]
    decision_rows = [[f"DEC-{decision_number(path):03d}", first_heading(path), status_from_markdown(path), pdf_path_breaks(path.relative_to(ROOT).as_posix())] for path in decision_files]
    story += [CondPageBreak(1.5*inch), Paragraph("13. DEC karar dizini — eksiksiz", h1), pdf_table(["ID", "Karar", "Durum", "Dosya"], decision_rows, [0.65*inch,3.1*inch,1.0*inch,2.25*inch], 5.5)]
    adr_rows = [[f"ADR-{adr_number(path):03d}", first_heading(path), pdf_path_breaks(path.relative_to(ROOT).as_posix())] for path in adr_files]
    story += [CondPageBreak(1.2*inch), Paragraph("14. ADR dizini — eksiksiz", h1), pdf_table(["ID", "Mimari karar", "Dosya"], adr_rows, [0.7*inch,3.9*inch,2.4*inch], 5.8)]
    rule_rows = [[rule["id"], rule["state"], rule["text"]] for rule in rules["rules"]]
    story += [CondPageBreak(1.2*inch), Paragraph(f"15. Kanonik kurallar — {rules['ruleCount']} kayıt", h1), pdf_table(["ID", "Durum", "Kural"], rule_rows, [0.65*inch,0.85*inch,5.5*inch], 6.0)]
    audit_rows = [[row["name"], row["type"].upper(), str(row["bytes"]), str(row.get("pages", "-")), "OK" if row["readable"] else "HATA"] for row in word_pdf_scan]
    story += [CondPageBreak(1.2*inch), Paragraph("16. Aktif repo Word/PDF tarihsel envanter denetimi", h1), pdf_table(["Dosya", "Tür", "Bayt", "Sayfa", "Okuma"], audit_rows, [4.1*inch,0.55*inch,0.8*inch,0.6*inch,0.95*inch], 6.1), p("Yeni sürüm tarihsel dosyaların yerine geçmez; yalnız aktif çalışma gerçeğini tek yerde toplar.")]
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    class FooterAfterContentDocument(SimpleDocTemplate):
        def afterPage(self):
            footer(self.canv, self)

    document = FooterAfterContentDocument(str(PDF_OUT), pagesize=letter, leftMargin=0.65*inch, rightMargin=0.65*inch, topMargin=0.68*inch, bottomMargin=0.78*inch, title="ParsYuva Aile Yaşam Merkezi Güncel Master Dokümantasyon", author="ParsYuva")
    document.build(story)


build_docx()
build_pdf()
print(SOURCE)
print(DOCX_OUT)
print(PDF_OUT)
print(f"source_sha256={sha256(SOURCE)}")
print(f"docx_sha256={sha256(DOCX_OUT)}")
print(f"pdf_sha256={sha256(PDF_OUT)}")
