from __future__ import annotations

import csv
import hashlib
import json
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader


PROJECT_ROOT = Path(r"C:\PPT\AYM")
REPO_ROOT = PROJECT_ROOT / "06_KOD" / "app"
JSON_OUT = REPO_ROOT / "artifacts" / "manifests" / "ALL_PROJECT_DOCUMENT_FORMAT_AUDIT.json"
MD_OUT = REPO_ROOT / "docs" / "current" / "12_TUM_BELGE_TURLERI_DENETIMI.md"
SELF_GENERATED_OUTPUTS = {
    JSON_OUT.resolve(),
    MD_OUT.resolve(),
    (REPO_ROOT / "docs" / "current" / "MASTER_PROJE_DOKUMANTASYONU_GUNCEL_17.08.2026_V1.docx").resolve(),
    (REPO_ROOT / "docs" / "current" / "MASTER_PROJE_DOKUMANTASYONU_GUNCEL_17.08.2026_V1.pdf").resolve(),
}

DOCUMENT_EXTENSIONS = {
    ".doc", ".docx", ".pdf", ".rtf", ".odt",
    ".xls", ".xlsx", ".ods", ".ppt", ".pptx", ".odp",
    ".md", ".txt", ".csv", ".json", ".yaml", ".yml", ".html", ".htm",
}
OFFICE_AND_PDF = {".doc", ".docx", ".pdf", ".rtf", ".odt", ".xls", ".xlsx", ".ods", ".ppt", ".pptx", ".odp"}
EXCLUDED_DIRECTORY_NAMES = {
    ".git", "node_modules", "tmp", ".tmp", ".tmp-runtime-dist", "dist", "coverage", "out", "release",
    "ESKI_TARIHLI_KAYITLAR",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def excluded(path: Path) -> bool:
    return any(part in EXCLUDED_DIRECTORY_NAMES for part in path.parts)


def classify(path: Path) -> str:
    relative = path.relative_to(PROJECT_ROOT)
    parts = {part.casefold() for part in relative.parts}
    name = path.name.upper()
    if "09_arsiv" in parts or "arşiv" in parts or "checkpoints" in parts:
        return "HISTORICAL"
    if "BUILD" in name and path.suffix.lower() in {".docx", ".pdf"}:
        return "HISTORICAL"
    if "docs" in parts and "current" in parts:
        return "ACTIVE_REFERENCE"
    if path.suffix.lower() in {".json", ".yaml", ".yml"}:
        return "MACHINE_READABLE_SOURCE_OR_EVIDENCE"
    return "SOURCE_OR_REFERENCE"


def read_text_strict(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def inspect(path: Path) -> dict:
    suffix = path.suffix.lower()
    row = {
        "path": path.relative_to(PROJECT_ROOT).as_posix(),
        "extension": suffix,
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
        "classification": classify(path),
        "readable": True,
    }
    try:
        if suffix == ".docx":
            document = Document(path)
            row["paragraphs"] = len(document.paragraphs)
            row["tables"] = len(document.tables)
            with ZipFile(path) as archive:
                row["packageEntryCount"] = len(archive.namelist())
        elif suffix == ".pdf":
            reader = PdfReader(str(path))
            row["pages"] = len(reader.pages)
            row["encrypted"] = bool(reader.is_encrypted)
            if not reader.is_encrypted:
                row["textCharacters"] = sum(len(page.extract_text() or "") for page in reader.pages)
        elif suffix == ".rtf":
            raw = path.read_bytes()
            text = None
            for encoding in ("utf-8-sig", "cp1254", "cp1252"):
                try:
                    text = raw.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if text is None or not text.lstrip().startswith("{\\rtf"):
                raise ValueError("RTF header or supported encoding missing")
            row["textCharacters"] = len(text)
        elif suffix == ".json":
            value = json.loads(read_text_strict(path))
            row["jsonRootType"] = type(value).__name__
        elif suffix == ".csv":
            with path.open("r", encoding="utf-8-sig", newline="") as stream:
                row["rows"] = sum(1 for _ in csv.reader(stream))
        elif suffix in {".md", ".txt", ".yaml", ".yml", ".html", ".htm"}:
            row["textCharacters"] = len(read_text_strict(path))
        elif suffix in {".xlsx", ".ods", ".pptx", ".odp", ".odt"}:
            with ZipFile(path) as archive:
                row["packageEntryCount"] = len(archive.namelist())
        elif suffix in {".doc", ".xls", ".ppt"}:
            signature = path.read_bytes()[:8]
            if signature != bytes.fromhex("D0CF11E0A1B11AE1"):
                raise ValueError("legacy OLE compound-document signature missing")
    except Exception as error:
        row["readable"] = False
        row["error"] = f"{type(error).__name__}: {error}"
    return row


previous_report = None
previous_historical_by_path = {}
if JSON_OUT.exists():
    previous_report = json.loads(JSON_OUT.read_text(encoding="utf-8"))
    previous_historical_by_path = {
        row["path"]: row
        for row in previous_report.get("files", [])
        if row.get("classification") == "HISTORICAL"
    }

paths = []
for path in PROJECT_ROOT.rglob("*"):
    if not path.is_file() or excluded(path):
        continue
    if path.resolve() in SELF_GENERATED_OUTPUTS:
        continue
    if path.suffix.lower() in DOCUMENT_EXTENSIONS:
        paths.append(path)
paths.sort(key=lambda value: value.relative_to(PROJECT_ROOT).as_posix().casefold())

files = []
for path in paths:
    relative = path.relative_to(PROJECT_ROOT).as_posix()
    classification = classify(path)
    if classification == "HISTORICAL" and previous_report is not None:
        previous = previous_historical_by_path.get(relative)
        if previous is not None:
            row = dict(previous)
            row["validationMode"] = "FROZEN_BASELINE_NOT_RECHECKED_DEC_252"
        else:
            row = {
                "path": relative,
                "extension": path.suffix.lower(),
                "bytes": path.stat().st_size,
                "sha256": None,
                "classification": "HISTORICAL",
                "readable": None,
                "validationMode": "NEW_HISTORICAL_ENTRY_NOT_REVIEWED_DEC_252",
            }
        files.append(row)
    else:
        row = inspect(path)
        row["validationMode"] = "CURRENT_ACTIVE_REVIEW"
        files.append(row)
extension_counts = dict(sorted(Counter(row["extension"] for row in files).items()))
classification_counts = dict(sorted(Counter(row["classification"] for row in files).items()))
folder_counts = dict(sorted(Counter(row["path"].split("/", 1)[0] for row in files).items()))
hash_counts = Counter(row["sha256"] for row in files if row.get("sha256"))
unreadable = [row for row in files if row.get("readable") is False]
not_rechecked = [row for row in files if row.get("readable") is None or row.get("validationMode", "").startswith("FROZEN_BASELINE")]
readable_count = sum(1 for row in files if row.get("readable") is True)
office = [row for row in files if row["extension"] in OFFICE_AND_PDF]

report = {
    "schemaVersion": 1,
    "id": "PPT-ALL-PROJECT-DOCUMENT-FORMAT-AUDIT-2026-08-17-V1",
    "root": str(PROJECT_ROOT),
    "generatedAt": datetime.now(timezone.utc).isoformat(),
    "exclusions": sorted(EXCLUDED_DIRECTORY_NAMES),
    "scopeNote": "Generated current-output folder and this audit's own generated files are excluded to prevent self-referential inventory; final DOCX/PDF/package contents are separately checksummed after packaging.",
    "documentFileCount": len(files),
    "readableCount": readable_count,
    "unreadableCount": len(unreadable),
    "frozenHistoricalNotRecheckedCount": len(not_rechecked),
    "officeAndPdfCount": len(office),
    "officeAndPdfReadableCount": sum(1 for row in office if row["readable"]),
    "uniqueContentHashCount": len(hash_counts),
    "duplicateFileCount": sum(count - 1 for count in hash_counts.values() if count > 1),
    "historicalContentReviewPolicy": {
        "decision": "DEC-252",
        "baselineCreated": previous_report is not None,
        "futureMode": "FROZEN_BASELINE_NOT_REOPENED_NOT_REVALIDATED",
        "activeAndNewDocumentsOnly": True
    },
    "extensionCounts": extension_counts,
    "classificationCounts": classification_counts,
    "topLevelFolderCounts": folder_counts,
    "absentOfficeFormats": [extension for extension in sorted(OFFICE_AND_PDF) if extension_counts.get(extension, 0) == 0],
    "unreadable": unreadable,
    "files": files,
}

JSON_OUT.parent.mkdir(parents=True, exist_ok=True)
JSON_OUT.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

lines = [
    "# Tüm Belge Türleri Denetimi",
    "",
    "- Sürüm: **GUNCEL-2026-08-17-V1**",
    f"- Kök: `{PROJECT_ROOT}`",
    f"- Denetlenen belge/config/metin dosyası: **{len(files)}**",
    f"- Okunabilir/önceki temelde okunabilirliği kanıtlı: **{readable_count}**",
    f"- Okunamayan/bozuk: **{len(unreadable)}**",
    f"- `DEC-252` gereği bu çalışmada yeniden açılmayan dondurulmuş tarihsel kayıt: **{len(not_rechecked)}**",
    f"- Office/RTF/PDF: **{len(office)}**; okunabilir **{sum(1 for row in office if row['readable'])}**",
    f"- Benzersiz içerik hash'i: **{len(hash_counts)}**; tekrar kopya: **{sum(count - 1 for count in hash_counts.values() if count > 1)}**",
    "",
    "> Her ana konu klasöründeki `ESKI_TARIHLI_KAYITLAR` dizini, dondurulmuş geçmiş kayıtların yeniden açılmasını veya güncellik denetimine alınmasını önlemek için kaynak taramasından çıkarılır. Güncel belgeler ilgili `00`-`11` konu klasörlerinde yerinde taranır.",
    "> Bu dosyanın ilk üretimi tarihsel kayıtlar için son içerik-okunabilirlik temelidir. Sonraki çalışmalarda tarihsel dosyalar önceki satırlarıyla taşınır; yeniden açılmaz, render edilmez veya semantik güncellik denetimine alınmaz.",
    "",
    "## Uzantı dağılımı",
    "",
    "| Uzantı | Dosya |",
    "|---|---:|",
]
for extension, count in extension_counts.items():
    lines.append(f"| `{extension}` | {count} |")
lines += ["", "## Üst klasör dağılımı", "", "| Klasör | Dosya |", "|---|---:|"]
for folder, count in folder_counts.items():
    lines.append(f"| `{folder}` | {count} |")
lines += [
    "",
    "## Sınıflandırma",
    "",
    "| Sınıf | Dosya |",
    "|---|---:|",
]
for classification, count in classification_counts.items():
    lines.append(f"| `{classification}` | {count} |")
lines += [
    "",
    "## Bulunmayan Office türleri",
    "",
    ", ".join(f"`{value}`" for value in report["absentOfficeFormats"]) or "Yok",
    "",
    "## Okuma sorunları",
    "",
]
if unreadable:
    for row in unreadable:
        lines.append(f"- `{row['path']}` — {row['error']}")
else:
    lines.append("- Yok. Denetlenen tüm dosyalar türüne uygun biçimde açıldı/ayrıştırıldı.")
lines += [
    "",
    "## Yetki ve tarih ilkesi",
    "",
    "- `09_ARSIV`, `Arşiv`, checkpoint ve geçmiş Build DOCX/PDF dosyaları tarihsel kanıttır; güncel otorite değildir.",
    "- Güncel iş gerçeği aktif config, `docs/current`, DEC/ADR/threat modelleri, kaynak kod ve testlerden üretilir.",
    "- Excel veya PowerPoint dosyası bulunmaması eksiklik olarak yorumlanmaz; bu projede tablo ve sunum otoritesi tanımlanmamıştır.",
    "- Tam dosya/yol/hash envanteri `artifacts/manifests/ALL_PROJECT_DOCUMENT_FORMAT_AUDIT.json` içindedir.",
    "",
]
MD_OUT.parent.mkdir(parents=True, exist_ok=True)
MD_OUT.write_text("\n".join(lines), encoding="utf-8")

print(f"files={len(files)} readable={readable_count} unreadable={len(unreadable)} frozen_historical_not_rechecked={len(not_rechecked)}")
print(f"office_and_pdf={len(office)} office_readable={sum(1 for row in office if row['readable'])}")
print(JSON_OUT)
print(MD_OUT)
