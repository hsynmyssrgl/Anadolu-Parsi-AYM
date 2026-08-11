from pathlib import Path
import json, re
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[1]
ledger=json.loads((ROOT/'config/master-build-ledger.json').read_text(encoding='utf-8'))
constitution=json.loads((ROOT/'config/project-constitution.json').read_text(encoding='utf-8'))
ui=json.loads((ROOT/'config/ui-visual-reference-manifest.json').read_text(encoding='utf-8'))
progress_model=json.loads((ROOT/'config/project-progress-model.json').read_text(encoding='utf-8'))
build=ledger['currentBuild']; version=ledger['currentVersion']; entry=next(x for x in ledger['builds'] if x['build']==build)
rules=ledger['projectRules']['versions'][-1]['rules']; open_items=[x for x in ledger['remainingWork'] if x['status'] in ('OPEN','IN_PROGRESS')]
out=ROOT/'docs/current'/f'MASTER_PROJECT_DOCUMENTATION_BUILD{build}.docx'

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_cell_text(cell,text,bold=False,color=None,size=9):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(str(text)); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.6);sec.bottom_margin=Inches(.6);sec.left_margin=Inches(.65);sec.right_margin=Inches(.65)
styles=doc.styles
styles['Normal'].font.name='Arial';styles['Normal'].font.size=Pt(9.5)
for s in ['Title','Heading 1','Heading 2','Heading 3']:
    styles[s].font.name='Arial'
styles['Title'].font.size=Pt(24);styles['Title'].font.bold=True
styles['Heading 1'].font.size=Pt(16);styles['Heading 1'].font.bold=True;styles['Heading 1'].font.color.rgb=RGBColor(124,77,32)
styles['Heading 2'].font.size=Pt(12);styles['Heading 2'].font.bold=True;styles['Heading 2'].font.color.rgb=RGBColor(79,44,14)

p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ANADOLU PARSI');r.bold=True;r.font.name='Arial';r.font.size=Pt(15);r.font.color.rgb=RGBColor(124,77,32)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Aile Yaşam Merkezi');r.bold=True;r.font.size=Pt(24)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(f'Master Proje Dokümantasyonu — Build {build}');r.font.size=Pt(15)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f'Sürüm {version} · Bronze RC2 Active Development · {entry["status"]}')

summary=doc.add_table(rows=0,cols=2);summary.alignment=WD_TABLE_ALIGNMENT.CENTER;summary.style='Table Grid'
for k,v in [('Üst marka','Panthera pardus tulliana'),('Uygulama','Anadolu Parsı Aile Yaşam Merkezi'),('Proje başlangıcı','20.07.2026'),('Build',build),('Kural seti',ledger['projectRules']['currentVersion']),('Kural sayısı',len(rules)),('Kural SHA-256',ledger['projectRules']['versions'][-1]['sha256']),('Aşama',ledger['currentStage'])]:
    cells=summary.add_row().cells;set_cell_text(cells[0],k,True,'FFFFFF');shade(cells[0],'7C4D20');set_cell_text(cells[1],v)

doc.add_paragraph('')
doc.add_heading('1. Proje Anayasası ve kaynak sınırı',1)
doc.add_paragraph('Bu proje için tek yetkili başlangıç 20 Temmuz 2026’dır. Bu tarihten önceki sohbet, dosya, karar veya proje bağlam kaynağı değildir. Güncel Proje Anayasası Ana Build Defteri içindeki kural setidir ve sessiz istisna kabul etmez.')
for text in [
    'Üst marka Panthera pardus tulliana; kullanıcıya görünen uygulama adı Anadolu Parsı Aile Yaşam Merkezi.',
    'Aktif kaynak, belge, görsel ve metadata doğal kişi/aile kimliği içermez; marka kimliği kullanılır.',
    'Production başlangıcı hazır kişi/aile/demo verisi içermez; veri yalnız kullanıcı işlemi veya kontrollü içe aktarma ile oluşur.',
    'Silver öncesi gerçek arayüz, UI Görsel Referans Manifestosu ve onaylı 20 Temmuz sonrası görsellerle doğrulanır.',
    'P0 yaşamsal API/adapter işleri önce gelir; P2 banka ve diğer kurum entegrasyonları kararlı üretimden yaklaşık 5-6 ay sonra değerlendirilir.'
]: doc.add_paragraph(text,style='List Bullet')

doc.add_heading('2. Güvenli ilk kullanım ve kullanıcı veri kasası',1)
for text in [
    'Windows ilk çalıştırma Anadolu Parsı marka tanıtımı, Türkçe sesli anlatım, altyazı/sessiz/geç seçenekleri ve ilk kurulum sihirbazıyla başlar.',
    'Yerel güçlü parola temel yöntemdir; Apple, Google ve Microsoft yalnız kimlik doğrulama sağlayıcısıdır ve uygulama içi yetki vermez.',
    'TOTP ve kurtarma kodları ilk güvenlik kurulumunun zorunlu parçasıdır; gerçek OIDC PASS ancak canlı sağlayıcı kayıtları ve Windows testleriyle verilir.',
    'Kullanıcı doğrulanmadan aile veritabanı açılmaz. Kalıcı ana veri AES-256-GCM kasada; veri anahtarı scrypt parola türetimi + Windows safeStorage/DPAPI cihaz bağıyla korunur.',
    'Kimliği doğrulanmış aktif oturumda SQLite ana veritabanı yalnız süreç belleğinde çalışır; kalıcı ana veri AES-256-GCM kasadır. Dosya görüntüsü gereken bounded staging Windows production’da EFS korumalı ve fail-closed’dur.',
    'Aktif oturum en fazla 30 saniyede bir şifreli kasaya checkpoint edilir; logout/timeout/quit son mühürleme ve staging temizliği yapar.',
    'Aynı Windows kullanıcısı yetkisindeki malware/yöneticiye karşı mutlak engel iddia edilmez. Build228, exact Build227 source ve gerçek Windows evidence SHA-256 bağında OPEN-021 ile OPEN-022 durumlarını resmen CLOSED yapmıştır; bu governance kararı Silver doğrulama sonuçlarını değiştirmez.'
]: doc.add_paragraph(text,style='List Bullet')

doc.add_heading('3. UI Görsel Referans Manifestosu',1)
doc.add_paragraph(f'Bağlayıcı görsel: {ui["image"]}. Ölçülebilir tipografi ve renk sözleşmesi config/ui-visual-reference-manifest.json ile kaynak CSS tarafından belirlenir.')
try:
    doc.add_picture(str(ROOT/ui['image']),width=Inches(7.1))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
except Exception:
    doc.add_paragraph('Görsel manifesto dosyası kaynak paketinde yer almaktadır.')

doc.add_heading('4. Kesin Proje Kuralları',1)
for rule in rules:
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(3);r=p.add_run(f'{rule["id"]} — ');r.bold=True;p.add_run(rule['text'])

doc.add_heading('5. Güncel devam noktası ve açık işler',1)
next_item=sorted(open_items,key=lambda x:x['order'])[0] if open_items else None
if next_item: doc.add_paragraph(f'Sıradaki iş: {next_item["id"]} — {next_item["title"]}. Planlanan Build: {next_item.get("plannedBuild","-")}.')
t=doc.add_table(rows=1,cols=5);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
for c,v in zip(t.rows[0].cells,['ID','İş','Kanal','Planlanan Build','Durum']): set_cell_text(c,v,True,'FFFFFF',8);shade(c,'7C4D20')
for item in sorted(open_items,key=lambda x:x['order']):
    cells=t.add_row().cells
    for c,v in zip(cells,[item['id'],item['title'],item['channel'],item.get('plannedBuild','-'),item['status']]):set_cell_text(c,v,size=7.7)

doc.add_heading('6. Build kapanış ve doğrulama yönetişimi',1)
for x in ['Project Provenance Gate','Version Sweep Gate','Personal Identity Sweep Gate','Production Clean Data Gate','Artifact Index Gate','Documentation Closure Gate','Project Progress Report Gate','Ana Build Defteri ve kural SHA-256 kabul kapısı','Sohbet bağlam kapasitesi / %90 hard-stop kapısı']:
    doc.add_paragraph(x,style='List Bullet')
doc.add_paragraph('Clean npm ci, tam root TypeScript, tüm testler, Electron production build, blocking smoke ve gerçek Windows installer yalnız gerçekten çalıştırıldığında PASS olabilir.')

doc.add_heading('7. Proje ilerleme modeli',1)
prog=entry.get('projectProgressAssessment')
if prog:
    for k,v in [('Kodlama tamamlanma',f'%{prog["codingCompletionPercent"]}'),('Kalan kodlama',f'%{prog["codingRemainingPercent"]}'),('Geçen süre',f'{prog["elapsedDays"]} gün'),('Tarihsel build hızı',f'{prog["historicalBuildsPerElapsedDay"]} build/gün'),('Tahmini Bronze Final',prog['estimatedBronzeFinalDate']),('Tahmini Silver',prog['estimatedSilverDate']),('Tahmini Gold/genel bitiş',prog['estimatedGoldDate']),('Güven',prog['confidence'])]:
        doc.add_paragraph(f'{k}: {v}')
else:
    doc.add_paragraph('Build kapanış ilerleme kaydı henüz oluşturulmadı. Model: config/project-progress-model.json.')

doc.add_heading('8. Yetkili aktif belgeler',1)
auth=(ROOT/'docs/11_DOCUMENT_AUTHORITY_MATRIX.md').read_text(encoding='utf-8')
for m in re.finditer(r'\| `([^`]+)`(?: ve `([^`]+)`)? \| ([^|]+) \| ([^|]+) \|',auth):
    paths=[m.group(1)]+([m.group(2)] if m.group(2) else [])
    doc.add_paragraph(' / '.join(paths)+' — '+m.group(3).strip(),style='List Bullet')

doc.add_heading(f'9. Build {build} karar ve kanıt dosyaları',1)
decision_files = (['docs/decisions/DEC-104-protected-side-artifact-encryption.md','docs/adr/ADR-087-protected-side-artifact-boundary.md','docs/decisions/DEC-105-pr171-atomic-work-segmentation.md','docs/adr/ADR-088-pr171-stepwise-validation-persistence.md','docs/security/PROTECTED_SIDE_ARTIFACTS_BUILD214.md','docs/decisions/DEC-106-windows-security-evidence-harness.md','docs/adr/ADR-089-real-windows-efs-dpapi-packaged-evidence.md','docs/security/WINDOWS_SECURITY_EVIDENCE_BUILD215.md','docs/decisions/DEC-107-windows-evidence-intake-source-binding.md','docs/adr/ADR-090-windows-evidence-intake-and-source-binding.md','docs/security/WINDOWS_EVIDENCE_INTAKE_BUILD216.md','docs/18_PROJECT_CONSTITUTION_V5.md','docs/18_PROJECT_CONSTITUTION_V5.json'] if build >= 216 else (['docs/decisions/DEC-104-protected-side-artifact-encryption.md','docs/adr/ADR-087-protected-side-artifact-boundary.md','docs/decisions/DEC-105-pr171-atomic-work-segmentation.md','docs/adr/ADR-088-pr171-stepwise-validation-persistence.md','docs/security/PROTECTED_SIDE_ARTIFACTS_BUILD214.md','docs/decisions/DEC-106-windows-security-evidence-harness.md','docs/adr/ADR-089-real-windows-efs-dpapi-packaged-evidence.md','docs/security/WINDOWS_SECURITY_EVIDENCE_BUILD215.md','docs/18_PROJECT_CONSTITUTION_V5.md','docs/18_PROJECT_CONSTITUTION_V5.json'] if build >= 215 else (['docs/decisions/DEC-104-protected-side-artifact-encryption.md','docs/adr/ADR-087-protected-side-artifact-boundary.md','docs/decisions/DEC-105-pr171-atomic-work-segmentation.md','docs/adr/ADR-088-pr171-stepwise-validation-persistence.md','docs/security/PROTECTED_SIDE_ARTIFACTS_BUILD214.md','docs/18_PROJECT_CONSTITUTION_V5.md','docs/18_PROJECT_CONSTITUTION_V5.json'] if build >= 214 else (['docs/decisions/DEC-103-memory-resident-user-data-session.md','docs/adr/ADR-086-memory-resident-sqlite-windows-efs-staging.md','docs/security/IN_USE_USER_DATA_PROTECTION_BUILD213.md'] if build >= 213 else (['docs/decisions/DEC-102-approved-ui-visual-baseline-correction.md','docs/adr/ADR-085-approved-ui-visual-baseline-hash-pinning.md','docs/ui/UI_VISUAL_REFERENCE_MANIFESTO.md'] if build >= 212 else (['docs/decisions/DEC-101-clean-install-external-access-handoff.md','docs/adr/ADR-084-clean-install-external-access-handoff.md','docs/NPM_CLEAN_INSTALL_BUILD211_HANDOFF.md'] if build >= 211 else ['docs/decisions/DEC-099-secure-onboarding-and-user-data-vault.md','docs/adr/ADR-082-secure-onboarding-user-data-vault.md']))))))
if build >= 217:
    decision_files += ['docs/decisions/DEC-108-open021-isolated-windows-closure-gate.md','docs/adr/ADR-091-open021-efs-only-real-windows-proof.md','docs/security/OPEN021_WINDOWS_CLOSURE_BUILD217.md']
if build >= 218:
    decision_files += ['docs/decisions/DEC-109-open022-isolated-windows-closure-gate.md','docs/adr/ADR-092-open022-dpapi-protected-side-artifact-proof.md','docs/security/OPEN022_WINDOWS_CLOSURE_BUILD218.md']
    decision_files += ['docs/decisions/DEC-110-unified-bronze-windows-security-closure.md','docs/adr/ADR-093-unified-bronze-windows-security-lifecycle.md','docs/security/BRONZE_WINDOWS_SECURITY_CLOSURE_BUILD219.md']
if build >= 220:
    decision_files += ['docs/decisions/DEC-111-build219-windows-failure-bootstrap-remediation.md','docs/adr/ADR-094-windows-packager-bootstrap-and-ps51-evidence-encoding.md','docs/security/BRONZE_WINDOWS_SECURITY_RETRY_BUILD220.md']
if build >= 221:
    decision_files += ['docs/decisions/DEC-112-build220-windows-failure-workspace-build-remediation.md','docs/adr/ADR-095-workspace-package-build-before-windows-package.md','docs/security/BRONZE_WINDOWS_SECURITY_RETRY_BUILD221.md']
if build >= 222:
    decision_files += ['docs/decisions/DEC-113-build221-windows-failure-preload-typescript-remediation.md','docs/adr/ADR-096-preload-global-lifecycle-typing.md','docs/security/BRONZE_WINDOWS_SECURITY_RETRY_BUILD222.md']
if build >= 223:
    decision_files += ['docs/decisions/DEC-114-build222-windows-failure-preload-cjs-graph-remediation.md','docs/adr/ADR-097-preload-commonjs-staging-graph.md','docs/security/BRONZE_WINDOWS_SECURITY_RETRY_BUILD223.md']
if build >= 224:
    decision_files += ['docs/decisions/DEC-115-build223-windows-failure-license-rtf-sync-remediation.md','docs/adr/ADR-098-deterministic-nsis-license-source-sync.md','docs/security/BRONZE_WINDOWS_SECURITY_RETRY_BUILD224.md']
if build >= 225:
    decision_files += ['docs/18_PROJECT_CONSTITUTION_V6.md','docs/18_PROJECT_CONSTITUTION_V6.json','docs/decisions/DEC-116-build224-windows-security-root-cause-remediation.md','docs/adr/ADR-099-fail-closed-windows-efs-safestorage-startup-evidence.md','docs/decisions/DEC-117-pr172-platform-actual-context-hard-stop.md','docs/adr/ADR-100-platform-actual-conversation-capacity-gate.md','docs/security/BRONZE_WINDOWS_SECURITY_RETRY_BUILD225.md']
if build >= 226:
    decision_files += ['docs/decisions/DEC-118-build225-fresh-profile-device-identity-initialization-order.md','docs/adr/ADR-101-protected-device-identity-before-device-bound-maintenance-restore.md','docs/security/FRESH_PROFILE_DEVICE_IDENTITY_INITIALIZATION_BUILD226.md']
if build >= 227:
    decision_files += ['docs/decisions/DEC-119-build227-four-proven-windows-root-causes.md','docs/adr/ADR-102-build227-windows-persistence-and-closure-remediation.md','docs/security/WINDOWS_ROOT_CAUSE_REMEDIATION_BUILD227.md']
if build >= 228:
    decision_files += ['docs/decisions/DEC-120-build228-open021-open022-official-closure.md','docs/adr/ADR-103-build227-evidence-bound-bronze-open-closure.md','docs/security/BRONZE_OPEN021_OPEN022_CLOSURE_BUILD228.md','config/bronze-open-closure-status.json']
for path in decision_files + [f'BUILD_STATUS_BRONZE_RC2_BUILD{build}.md',f'RELEASE_NOTES_BRONZE_RC2_BUILD{build}.md',f'BUILD{build}_ARCHITECTURE_VALIDATION_REPORT.md',f'BUILD{build}_DELIVERY_VALIDATION_REPORT.md','PROJECT_ARTIFACT_INDEX.md','PROJECT_ARTIFACT_INDEX.json']:
    doc.add_paragraph(path,style='List Bullet')

doc.add_paragraph('')
p=doc.add_paragraph('Bu belge güncel kaynak ağacından Build kapanış sürecinde otomatik üretilir. Tarihsel build belgeleri kendi sürümlerini korur; aktif davranış için Ana Build Defteri ve güncel Proje Anayasası üstündür.')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
out.parent.mkdir(parents=True,exist_ok=True);doc.save(out);print(out)
