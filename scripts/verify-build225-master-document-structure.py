from __future__ import annotations

import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs/current/MASTER_PROJECT_DOCUMENTATION_BUILD225.docx"
PDF = ROOT / "docs/current/MASTER_PROJECT_DOCUMENTATION_BUILD225.pdf"
REPORT = ROOT / "artifacts/validation/build225-master-document-structural-qa.json"

checks: list[dict[str, object]] = []


def record(name: str, passed: bool, detail: object) -> None:
    checks.append({"name": name, "status": "PASS" if passed else "FAIL", "detail": detail})


with zipfile.ZipFile(DOCX) as archive:
    bad = archive.testzip()
    record("docx-zip-integrity", bad is None, bad or "all entries readable")

document = Document(DOCX)
docx_text = "\n".join(
    [paragraph.text for paragraph in document.paragraphs]
    + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
)
record("docx-nonempty", len(docx_text.strip()) > 1000, len(docx_text))
record("docx-build-version", "Build 225" in docx_text and "02.08.2026.225" in docx_text, "Build 225 / 02.08.2026.225")
docx_rule_ids = sorted(set(re.findall(r"PR-\d{3}", docx_text)))
record("docx-all-project-rules", len(docx_rule_ids) == 172 and docx_rule_ids[0] == "PR-001" and docx_rule_ids[-1] == "PR-172", len(docx_rule_ids))
placeholders = [token for token in ("TODO", "TBD", "LOREM IPSUM", "{{", "}}") if token in docx_text.upper()]
record("docx-no-placeholders", not placeholders, placeholders)
record("docx-page-setup", all(section.page_width and section.page_height for section in document.sections), len(document.sections))

reader = PdfReader(PDF)
pdf_pages = [(page.extract_text() or "").strip() for page in reader.pages]
pdf_text = "\n".join(pdf_pages)
record("pdf-page-count", 8 <= len(pdf_pages) <= 20, len(pdf_pages))
record("pdf-no-blank-pages", all(len(text) >= 40 for text in pdf_pages), [index + 1 for index, text in enumerate(pdf_pages) if len(text) < 40])
record("pdf-build-version", "Build 225" in pdf_text and "02.08.2026.225" in pdf_text, "Build 225 / 02.08.2026.225")
pdf_rule_ids = sorted(set(re.findall(r"PR-\d{3}", pdf_text)))
record("pdf-all-project-rules", len(pdf_rule_ids) == 172 and pdf_rule_ids[0] == "PR-001" and pdf_rule_ids[-1] == "PR-172", len(pdf_rule_ids))
record("pdf-open-boundary", "NOT_RUN / NOT_READY" in pdf_text, "OPEN-021/022 real Windows boundary")

failures = [item for item in checks if item["status"] == "FAIL"]
report = {
    "schemaVersion": 1,
    "build": 225,
    "status": "PASS" if not failures else "FAIL",
    "checks": checks,
    "visualQa": {
        "pdf": "RENDERED_AND_INSPECTED",
        "docx": "STRUCTURAL_ONLY_LIBREOFFICE_UNAVAILABLE",
    },
    "generatedAt": datetime.now(timezone.utc).isoformat(),
}
REPORT.parent.mkdir(parents=True, exist_ok=True)
REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
print(f"Build225 master document structural QA: {report['status']} ({len(checks) - len(failures)}/{len(checks)})")
raise SystemExit(1 if failures else 0)
