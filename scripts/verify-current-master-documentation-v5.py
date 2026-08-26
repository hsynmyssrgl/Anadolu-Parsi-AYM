from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
ACTIVE = json.loads((ROOT / "config/active-document-set.json").read_text(encoding="utf-8"))
MASTER = ACTIVE["currentMasterDocumentation"]
EXPECTED_VERSION = str(MASTER.get("version", ""))
DOCX_PATH = ROOT / MASTER["docx"]
PDF_PATH = ROOT / MASTER["pdf"]
RULES = json.loads((ROOT / "config/canonical-rule-registry.json").read_text(encoding="utf-8"))["rules"]
LEDGER_DECISIONS = json.loads((ROOT / "config/user-decision-ledger.json").read_text(encoding="utf-8"))["decisions"]
DECISION_IDS = sorted(
    match.group(1)
    for path in (ROOT / "docs/decisions").glob("DEC-*.md")
    if (match := re.match(r"^(DEC-\d+)", path.name))
)
ADR_IDS = sorted(path.stem.split("-", 2)[0] + "-" + path.stem.split("-", 2)[1]
                 for path in (ROOT / "docs/adr").glob("ADR-*.md"))
ADR_NUMBERS = sorted(int(identifier.removeprefix("ADR-")) for identifier in ADR_IDS)
failures: list[str] = []


def check(condition: bool, message: str) -> None:
    if not condition:
        failures.append(message)


check(re.fullmatch(r"GUNCEL-\d{4}-\d{2}-\d{2}-V5", EXPECTED_VERSION) is not None,
      "current master version is not a dated V5 marker")
check(MASTER.get("asOf") == EXPECTED_VERSION.removeprefix("GUNCEL-").removesuffix("-V5"),
      "current master asOf does not match its version marker")
check(MASTER.get("status") == "ACTIVE_CURRENT_MASTER_REFERENCE", "current master status is not active")
check(MASTER.get("historicalBuildArtifactsImmutable") is True, "historical build immutability is not retained")
check(DOCX_PATH.is_file() and DOCX_PATH.stat().st_size > 0, "current master DOCX is missing or empty")
check(PDF_PATH.is_file() and PDF_PATH.stat().st_size > 0, "current master PDF is missing or empty")

docx_text = ""
docx_tables = 0
if DOCX_PATH.is_file():
    document = Document(DOCX_PATH)
    docx_tables = len(document.tables)
    docx_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    check(docx_tables >= 10, f"current master DOCX table count is unexpectedly low: {docx_tables}")

pdf_text = ""
pdf_pages = 0
if PDF_PATH.is_file():
    reader = PdfReader(str(PDF_PATH))
    check(not reader.is_encrypted, "current master PDF is encrypted")
    pdf_pages = len(reader.pages)
    page_texts = [page.extract_text() or "" for page in reader.pages]
    pdf_text = "\n".join(page_texts)
    check(pdf_pages >= 20, f"current master PDF page count is unexpectedly low: {pdf_pages}")
    for page_number, text in enumerate(page_texts, start=1):
        check(len(re.findall(rf"\bSayfa\s+{page_number}\b", text)) == 1,
              f"PDF page {page_number} does not contain one exact page marker")

rule_ids = sorted(str(item["id"]) for item in RULES)
decision_ids = DECISION_IDS
ledger_decision_ids = sorted(str(item["id"]) for item in LEDGER_DECISIONS)
check(set(ledger_decision_ids).issubset(decision_ids), "decision ledger contains an ID without a decision document")
check(ADR_NUMBERS == list(range(1, max(ADR_NUMBERS, default=0) + 1)),
      "ADR source numbering is not contiguous from ADR-001")
for label, identifiers in (("rule", rule_ids), ("decision", decision_ids), ("ADR", ADR_IDS)):
    for identifier in identifiers:
        check(identifier in docx_text, f"DOCX missing {label} identifier: {identifier}")
        check(identifier in pdf_text, f"PDF missing {label} identifier: {identifier}")

check(EXPECTED_VERSION in docx_text, "DOCX version marker is missing")
check(EXPECTED_VERSION in pdf_text, "PDF version marker is missing")
if failures:
    raise SystemExit("Current master V5 verification failed:\n" + "\n".join(failures))
print(
    f"Current master V5: PASS ({len(rule_ids)} rules / {len(decision_ids)} decisions / "
    f"{len(ADR_IDS)} ADR / {docx_tables} DOCX tables / {pdf_pages} PDF pages)."
)
